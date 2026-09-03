"""
Travel Itinerary Agent — Streamlit Web UI
Run:  streamlit run app.py

"""

import io
import json
import os
import re
import shutil
from datetime import datetime, timedelta
from pathlib import Path

# ── Proxy (corporate network only — comment out or unset if not needed) ───────
# If HTTPS_PROXY env var is already set in the shell, that value is used.
# On a home/mac machine with no proxy, leave these unset or set TRAVEL_PROXY="" to skip.
_proxy = os.environ.get("TRAVEL_PROXY", "")   # set to "" to disable
if _proxy:
    os.environ.setdefault("HTTPS_PROXY", _proxy)
    os.environ.setdefault("HTTP_PROXY",  _proxy)

import streamlit as st

# ── Paths — writable dirs use /tmp/ for Streamlit Community Cloud ──────────────
_APP_DIR     = Path(__file__).parent
DOCS_DIR     = _APP_DIR / "docs"
CHROMA_DIR   = Path("/tmp/.chroma_db")
HISTORY_DIR  = Path("/tmp/.chat_history")
TEMPLATE_DIR = Path("/tmp/template")
RETENTION    = 5   # days

# ── Knowledge source selection (Existing Knowledge Base vs Master Travel Plans) ─
# Two independent, isolated knowledge sources the user can pick from in the
# sidebar. "existing" is the original PDF/DOCX + ChromaDB RAG pipeline
# (build_vectorstore/get_answer, unchanged). "master_travel_plans" is the new
# Excel-based structured source (docs/All_Plans_of_KUKUTRIP_Master.xlsx),
# implemented entirely separately below so the existing source's behavior,
# indexing, and data are never touched by this feature.
KB_SOURCE_EXISTING = "Existing Knowledge Base"
KB_SOURCE_MASTER_PLANS = "KukuTrip Master Travel Plans"
MASTER_PLANS_XLSX = DOCS_DIR / "All_Plans_of_KUKUTRIP_Master.xlsx"

# ── Company branding (constant across all itineraries) ────────────────────────
COMPANY_NAME  = "Ankur Sharma"
COMPANY_PHONE = "+918929838899"
COMPANY_WEB   = "www.kukutrip.com"
COMPANY_EMAIL = "info@kukutrip.com"

# ── PDF Theme System ────────────────────────────────────────────────────────
# Centralized theme configuration — the SINGLE source of truth for every
# color used anywhere in PDF generation. No hex codes should be hardcoded
# inside generate_pdf()/generate_pdf_editorial(); they must all be looked up
# from the selected theme dict below via semantic keys (theme["primary"],
# theme["accent"], etc.). This is what lets a new theme be added later by
# just adding a new dict entry, without touching the PDF templates at all.
#
# Each theme declares a "layout":
#   "classic"   → the original KukuTrip design (gradient hero, rounded
#                 cards, pill badges) — see generate_pdf(). All 7 color
#                 themes below share this exact layout; only the palette
#                 changes.
#   "editorial" → the alternate document-style layout inspired by the
#                 reference Georgia_Tbilisi_Grand_Explorer_Itinerary PDF —
#                 see generate_pdf_editorial(). Cleaner typography, no
#                 hero photo, large day numbers, thin rule dividers.
#
# Semantic keys used by both layouts:
#   primary        — main heading / accent color
#   primary_dark   — darker shade, used for gradient starts / hero overlay
#   primary_light  — lighter shade, used for secondary gradient stops
#   accent         — secondary highlight color (gold for Navy & Gold, a
#                    restrained bronze for Modern Editorial, etc.) — used
#                    for price/footer gradients and day-number accents
#   pill_bg/pill_border/pill_text — metadata pill badge styling (classic)
#   bg_light       — subtle tinted background for cards/sections
#   hotel_row_alt  — alternating hotel-table row background
#   border         — default border color for cards/tables/dividers
#   text           — primary body text color
#   muted_text     — secondary/caption text color
#   white          — always #FFFFFF, kept for readability of the mapping
PDF_THEMES = {
    "Maroon Red": {
        "layout": "classic",
        "primary": "#9E1B1B", "primary_dark": "#7A1515",
        "primary_light": "#C62828", "accent": "#C62828",
        "pill_bg": "#FBECEC", "pill_border": "#E8DAD6", "pill_text": "#8F1C1C",
        "bg_light": "#FCF8F6", "hotel_row_alt": "#FBF0F0", "border": "#E8DAD6",
        "text": "#333333", "muted_text": "#444444", "white": "#FFFFFF",
    },
    "Modern Editorial": {
        "layout": "editorial",
        "primary": "#20242B", "primary_dark": "#0F1216",
        "primary_light": "#3A404B", "accent": "#A8763B",
        "pill_bg": "#FAFAF9", "pill_border": "#E5E7EB", "pill_text": "#20242B",
        "bg_light": "#FAFAF9", "hotel_row_alt": "#F5F5F4", "border": "#E5E7EB",
        "text": "#2B2B2B", "muted_text": "#6B7280", "white": "#FFFFFF",
    },
    "Modern Blue": {
        "layout": "classic",
        "primary": "#1A4D8F", "primary_dark": "#0F3466",
        "primary_light": "#2E6FC4", "accent": "#2E6FC4",
        "pill_bg": "#EAF1FB", "pill_border": "#D2E1F2", "pill_text": "#1A4D8F",
        "bg_light": "#F6FAFE", "hotel_row_alt": "#EEF4FC", "border": "#D2E1F2",
        "text": "#2B2B2B", "muted_text": "#4A4A4A", "white": "#FFFFFF",
    },
    "Emerald Green": {
        "layout": "classic",
        "primary": "#1B7A4A", "primary_dark": "#125A37",
        "primary_light": "#2E9E63", "accent": "#2E9E63",
        "pill_bg": "#E8F7EE", "pill_border": "#CFEBDA", "pill_text": "#1B7A4A",
        "bg_light": "#F5FBF7", "hotel_row_alt": "#EAF8EF", "border": "#CFEBDA",
        "text": "#2B2B2B", "muted_text": "#4A4A4A", "white": "#FFFFFF",
    },
    "Navy & Gold": {
        "layout": "classic",
        "primary": "#132A52", "primary_dark": "#0A1830",
        "primary_light": "#24406E", "accent": "#B8912F",
        "pill_bg": "#EFF2F8", "pill_border": "#DCE2EE", "pill_text": "#132A52",
        "bg_light": "#F7F8FB", "hotel_row_alt": "#EDF0F6", "border": "#DCE2EE",
        "text": "#242424", "muted_text": "#4A4A4A", "white": "#FFFFFF",
    },
    "Teal": {
        "layout": "classic",
        "primary": "#0F7A76", "primary_dark": "#0A5A57",
        "primary_light": "#17A39D", "accent": "#17A39D",
        "pill_bg": "#E7F6F5", "pill_border": "#CDEAE8", "pill_text": "#0F7A76",
        "bg_light": "#F4FBFB", "hotel_row_alt": "#E9F7F6", "border": "#CDEAE8",
        "text": "#2B2B2B", "muted_text": "#4A4A4A", "white": "#FFFFFF",
    },
    "Purple": {
        "layout": "classic",
        "primary": "#5B2A86", "primary_dark": "#3E1C5E",
        "primary_light": "#7C4AAE", "accent": "#7C4AAE",
        "pill_bg": "#F2ECF8", "pill_border": "#E1D3EF", "pill_text": "#5B2A86",
        "bg_light": "#FAF8FC", "hotel_row_alt": "#F3EDF9", "border": "#E1D3EF",
        "text": "#2B2B2B", "muted_text": "#4A4A4A", "white": "#FFFFFF",
    },
    "Minimal Black & White": {
        "layout": "classic",
        "primary": "#1A1A1A", "primary_dark": "#000000",
        "primary_light": "#4D4D4D", "accent": "#4D4D4D",
        "pill_bg": "#F2F2F2", "pill_border": "#DDDDDD", "pill_text": "#1A1A1A",
        "bg_light": "#FAFAFA", "hotel_row_alt": "#F2F2F2", "border": "#DDDDDD",
        "text": "#1A1A1A", "muted_text": "#555555", "white": "#FFFFFF",
    },
}
DEFAULT_PDF_THEME = "Maroon Red"

st.set_page_config(page_title="✈ Travel Itinerary Agent", page_icon="✈", layout="wide")


# ── Session / history helpers ─────────────────────────────────────────────────

def _session_ts(stem: str) -> datetime:
    return datetime.strptime(stem, "%Y-%m-%d_%H-%M-%S")

def purge_old_sessions():
    if not HISTORY_DIR.exists():
        return
    cutoff = datetime.now() - timedelta(days=RETENTION)
    for f in HISTORY_DIR.glob("*.json"):
        try:
            if _session_ts(f.stem) < cutoff:
                f.unlink()
        except Exception:
            pass

def list_sessions():
    if not HISTORY_DIR.exists():
        return []
    return sorted(HISTORY_DIR.glob("*.json"), reverse=True)

def save_session(sid: str, msgs: list):
    HISTORY_DIR.mkdir(exist_ok=True)
    (HISTORY_DIR / f"{sid}.json").write_text(
        json.dumps({"session_id": sid, "messages": msgs}, indent=2)
    )

def load_session(path: Path) -> list:
    return json.loads(path.read_text())["messages"]

def new_sid() -> str:
    return datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

purge_old_sessions()

if "session_id" not in st.session_state:
    st.session_state.session_id = new_sid()

# ── PDF generation — Italy template pixel-perfect replication ─────────────────

HEADER_CACHE = Path("/tmp/.header_cache")
HEADER_CACHE.mkdir(exist_ok=True)

def _wiki_image(search: str, proxies: dict, headers: dict) -> str | None:
    """Fetch a Wikipedia article's originalimage URL, skipping flags/SVGs/maps."""
    import requests
    try:
        api = f"https://en.wikipedia.org/api/rest_v1/page/summary/{search.replace(' ', '_')}"
        r = requests.get(api, proxies=proxies, timeout=12, headers=headers)
        if r.status_code != 200:
            return None
        data = r.json()
        img_url = (data.get("originalimage") or data.get("thumbnail") or {}).get("source", "")
        if not img_url:
            return None
        skip_words = ("Flag", "Map", "Coat_of", "Emblem", "Seal_of", "Logo")
        if img_url.lower().endswith(".svg") or any(w in img_url for w in skip_words):
            return None
        return img_url
    except Exception:
        return None


def _wikimedia_search_image(query: str, proxies: dict, headers: dict) -> str | None:
    """Search Wikimedia Commons for a landscape image matching the query."""
    import requests
    try:
        params = {
            "action": "query",
            "list": "search",
            "srsearch": f"{query} landscape",
            "srnamespace": "6",   # File: namespace
            "srlimit": "5",
            "format": "json",
        }
        r = requests.get("https://commons.wikimedia.org/w/api.php",
                         params=params, proxies=proxies, timeout=12, headers=headers)
        if r.status_code != 200:
            return None
        results = r.json().get("query", {}).get("search", [])
        for res in results:
            title = res.get("title", "")
            if not title.lower().endswith((".jpg", ".jpeg", ".png")):
                continue
            # Get direct URL for this file
            info_r = requests.get(
                "https://commons.wikimedia.org/w/api.php",
                params={"action": "query", "titles": title,
                        "prop": "imageinfo", "iiprop": "url", "format": "json"},
                proxies=proxies, timeout=10, headers=headers)
            if info_r.status_code == 200:
                pages = info_r.json().get("query", {}).get("pages", {})
                for page in pages.values():
                    url = (page.get("imageinfo") or [{}])[0].get("url", "")
                    skip_words = ("Flag", "Map", "Coat_of", "Emblem", "Seal_of", "Logo")
                    if url and not url.lower().endswith(".svg") and not any(w in url for w in skip_words):
                        return url
    except Exception:
        pass
    return None


def fetch_destination_image(keyword: str, _dbg: list | None = None) -> Path | None:
    """
    Dynamically fetch a scenic hero image for the destination.

    Strategy (in priority order):
    1. Specific landmark/location Wikipedia article (e.g. "Dolomites")
    2. Tourism_in_{Country} article
    3. Geography_of_{Country} article
    4. Country Wikipedia article
    5. Wikimedia Commons search for the keyword
    6. picsum.photos seeded fallback
    """
    import hashlib, requests
    slug = hashlib.md5(keyword.lower().encode()).hexdigest()[:12]
    cached = HEADER_CACHE / f"{slug}.jpg"
    if cached.exists():
        return cached

    _p = os.environ.get("HTTPS_PROXY") or os.environ.get("HTTP_PROXY") or ""
    proxies = {"https": _p, "http": _p} if _p else {}
    headers_h = {"User-Agent": "TravelItineraryAgent/1.0 (travel-pdf-generator; free use)"}

    def _log(msg):
        if _dbg is not None:
            _dbg.append(msg)

    # Build search candidates — keyword may be "Dolomites Italy" or "Kyoto" or "Tbilisi"
    words = keyword.strip().split()
    primary = words[0].title()
    country = words[-1].title() if len(words) > 1 else primary

    candidates = [
        keyword.replace(" ", "_").title(),
        primary,
        f"Tourism_in_{country}",
        f"Tourism_in_{primary}",
        f"Geography_of_{country}",
        f"{country}",
    ]

    img_url = None
    for search in candidates:
        img_url = _wiki_image(search, proxies, headers_h)
        if img_url:
            _log(f"✅ Wikipedia image found via '{search}'")
            break
    if not img_url:
        _log(f"⚠️ Wikipedia: no image found for any candidate")

    # Wikimedia Commons full-text search
    if not img_url:
        img_url = _wikimedia_search_image(keyword, proxies, headers_h)
        if img_url:
            _log(f"✅ Wikimedia Commons image found")
        else:
            _log(f"⚠️ Wikimedia Commons: no image found")

    if not img_url:
        # Last fallback: picsum.photos — reliable, free, no API key, seeded by keyword
        try:
            seed = int(hashlib.md5(keyword.lower().encode()).hexdigest()[:8], 16) % 1000000
            picsum_url = f"https://picsum.photos/seed/{seed}/1600/900"
            img_r = requests.get(picsum_url, proxies=proxies, timeout=20,
                                 headers=headers_h, allow_redirects=True)
            if img_r.status_code == 200 and len(img_r.content) > 15000:
                from PIL import Image as PILImage
                from io import BytesIO as BytesIO2
                pil = PILImage.open(BytesIO2(img_r.content)).convert("RGB")
                pil.save(str(cached), "JPEG", quality=92, optimize=True)
                _log(f"✅ picsum.photos fallback used")
                return cached
            else:
                _log(f"⚠️ picsum.photos: status={img_r.status_code}, size={len(img_r.content)}")
        except Exception as e:
            _log(f"❌ picsum.photos error: {e}")
        return None

    # Download and convert to landscape JPEG — retry up to 3 times
    for attempt in range(1, 4):
        try:
            img_r = requests.get(img_url, proxies=proxies, timeout=45, headers=headers_h)
            if img_r.status_code == 200 and len(img_r.content) > 15000:
                from PIL import Image as PILImage
                from io import BytesIO as BytesIO2
                pil = PILImage.open(BytesIO2(img_r.content)).convert("RGB")
                # Crop to 16:9 if portrait
                w, h = pil.size
                if h > w:
                    new_h = int(w * 9 / 16)
                    top = max(0, (h - new_h) // 4)
                    pil = pil.crop((0, top, w, top + new_h))
                # Resize to max 1200×675 — keeps file small for base64 embed in PDF
                pil.thumbnail((1200, 675), PILImage.LANCZOS)
                pil.save(str(cached), "JPEG", quality=85, optimize=True)
                _log(f"✅ Image downloaded and saved ({pil.width}x{pil.height}px, attempt {attempt})")
                return cached
            else:
                _log(f"⚠️ Attempt {attempt}: status={img_r.status_code}, size={len(img_r.content)}")
                if img_r.status_code in (429, 503) and attempt < 3:
                    import time; time.sleep(2 * attempt)
                else:
                    break
        except Exception as e:
            _log(f"❌ Attempt {attempt} error: {e}")
            if attempt < 3:
                import time; time.sleep(2)

    # Final fallback after Wikipedia download failure: picsum
    try:
        seed = int(hashlib.md5(keyword.lower().encode()).hexdigest()[:8], 16) % 1000000
        picsum_url = f"https://picsum.photos/seed/{seed}/1600/900"
        img_r = requests.get(picsum_url, proxies=proxies, timeout=20,
                             headers=headers_h, allow_redirects=True)
        if img_r.status_code == 200 and len(img_r.content) > 15000:
            from PIL import Image as PILImage
            from io import BytesIO as BytesIO2
            pil = PILImage.open(BytesIO2(img_r.content)).convert("RGB")
            pil.thumbnail((1200, 675), PILImage.LANCZOS)
            pil.save(str(cached), "JPEG", quality=85, optimize=True)
            _log(f"✅ picsum.photos used after Wikipedia download failure")
            return cached
        else:
            _log(f"⚠️ picsum fallback also failed: {img_r.status_code}")
    except Exception as e:
        _log(f"❌ picsum fallback error: {e}")
    return None

def _h(text: str) -> str:
    """HTML-escape a string."""
    return (str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _extract_amount_value(amount_str: str) -> float | None:
    """Extract the first numeric value from a price string, ignoring currency symbols/commas."""
    if not amount_str:
        return None
    m = re.search(r'[\d,]+(?:\.\d+)?', str(amount_str))
    if not m:
        return None
    try:
        return float(m.group().replace(",", ""))
    except Exception:
        return None


def _extract_persons_count(persons_str: str) -> int:
    """Extract number of adults/persons from a string like '4 Adults · 1 Room'."""
    try:
        m = re.search(r'(\d+)\s*[Aa]dult', persons_str or "")
        if m:
            return max(1, int(m.group(1)))
    except Exception:
        pass
    return 1


def _format_nights_days(nights_str: str, days_list=None) -> str:
    """Normalize any nights/days representation into 'N Nights / N Days'.
    Handles bare numbers ('9'), 'N Nights' only, 'N Days' only, or an
    already-correct 'N Nights / N Days' string. Falls back to counting
    the days[] array length when no days count can be derived otherwise."""
    import re as _re
    s = str(nights_str or "").strip()
    if not s:
        if days_list:
            d = len(days_list)
            if d > 0:
                return f"{max(0, d-1)} Nights / {d} Days"
        return ""
    m_n = _re.search(r'(\d+)\s*[Nn]ight', s)
    m_d = _re.search(r'(\d+)\s*[Dd]ay', s)
    if m_n and m_d:
        return f"{m_n.group(1)} Nights / {m_d.group(1)} Days"
    if m_n:
        n = int(m_n.group(1))
        return f"{n} Nights / {n+1} Days"
    if m_d:
        d = int(m_d.group(1))
        return f"{max(0, d-1)} Nights / {d} Days"
    m_bare = _re.search(r'^(\d+)$', s)
    if m_bare:
        n = int(m_bare.group(1))
        return f"{n} Nights / {n+1} Days"
    return s


def _is_usd_amount(amount_str: str) -> bool:
    """Detect whether a price string is denominated in USD."""
    return bool(re.search(r'USD|US\$|\$', str(amount_str or ""), re.I))


def _is_per_person_amount(amount_str: str) -> bool:
    """Detect whether a price string is denominated per person (vs. total package)."""
    return bool(re.search(r'per\s*person|per\s*pax|/\s*person|/\s*pax|\bpp\b',
                          str(amount_str or ""), re.I))



def _format_inr(amount: float) -> str:
    """Format a number using Indian digit grouping, e.g. 400000 -> '4,00,000'."""
    n = int(round(amount))
    sign = "-" if n < 0 else ""
    n = abs(n)
    s = str(n)
    if len(s) <= 3:
        return sign + s
    last3 = s[-3:]
    rest = s[:-3]
    parts = []
    while len(rest) > 2:
        parts.insert(0, rest[-2:])
        rest = rest[:-2]
    if rest:
        parts.insert(0, rest)
    return sign + ",".join(parts) + "," + last3


def _parse_markup(markup_raw: str, base_amount: float) -> float:
    """Parse a markup string (percentage or fixed amount) and return the markup value in INR."""
    markup_raw = (markup_raw or "").strip()
    if "%" in markup_raw:
        m = re.search(r'[\d.]+', markup_raw)
        pct = float(m.group()) if m else 0.0
        return base_amount * (pct / 100.0)
    val = _extract_amount_value(markup_raw)
    return val or 0.0




def _hex_to_rgba(hex_color: str, alpha: float) -> str:
    """Convert a '#RRGGBB' hex color to an 'rgba(r,g,b,a)' CSS string."""
    h = (hex_color or "#000000").lstrip('#')
    if len(h) != 6:
        h = "000000"
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return f"rgba({r},{g},{b},{alpha})"


def generate_pdf(title: str, content: str, meta: dict,
                 header_img_path: Path | None = None,
                 theme: dict | None = None) -> bytes:
    """
    Render the CLASSIC KukuTrip HTML/CSS travel itinerary template via
    WeasyPrint: rounded cards, gradient hero, pill badges, day-number
    tiles, price card, footer. This is the "classic" layout shared by
    every color theme except "Modern Editorial" (which uses
    generate_pdf_editorial() instead) — only the `theme` dict's colors
    change between calls; the markup/structure never changes.
    Uses local DejaVu Sans fonts — no internet required.
    """
    import base64
    from weasyprint import HTML as WH

    theme = theme or PDF_THEMES[DEFAULT_PDF_THEME]


    # ── Embed company logo (coollogo.png) ─────────────────────────────────────
    logo_tag = ""
    logo_path = _APP_DIR / "template" / "coollogo.png"
    if logo_path.exists():
        logo_b64 = base64.b64encode(logo_path.read_bytes()).decode()
        logo_tag = (
            f'<img class="hero-logo" src="data:image/png;base64,{logo_b64}" />'
        )

    # ── Embed hero image as base64 data-URI ───────────────────────────────────
    if header_img_path and header_img_path.exists():
        img_b64 = base64.b64encode(header_img_path.read_bytes()).decode()
        # Dark gradient overlay ensures text is always readable over the photo
        hero_bg = (
            "linear-gradient(180deg,rgba(30,10,40,0.68) 0%,"
            "rgba(100,20,40,0.52) 50%,rgba(30,10,40,0.72) 100%),"
            f"url('data:image/jpeg;base64,{img_b64}') center/cover no-repeat"
        )
    else:
        hero_bg = (
            f"linear-gradient(180deg,{theme['primary_dark']} 0%,"
            f"{theme['primary']} 45%,{theme['primary_light']} 100%)"
        )


    pkg_name  = _h(meta.get("package_name", meta.get("destination", title)).upper())
    route_txt = _h(meta.get("route", ""))
    dates_txt = _h(meta.get("dates", ""))
    nights_txt = _h(_format_nights_days(meta.get("nights", ""), meta.get("days")))

    # ── Hero dates line: prefer dates, else show nights ────────────────────────
    hero_dates = dates_txt or nights_txt

    # ── Metadata pills — centered group ───────────────────────────────────────
    pill_items = [
        meta.get("dates", ""), _format_nights_days(meta.get("nights", ""), meta.get("days")),
        meta.get("persons", ""), meta.get("transport", ""),
    ]
    pills_html = "".join(
        f'<span class="pill">{_h(p)}</span>'
        for p in pill_items if p
    )

    # ── Final customer-facing price ─────────────────────────────────────────
    # If a finalized price has been computed via the pricing workflow
    # (_final_price_label / _final_price_value), use that exclusively.
    # Never fall back to showing raw/base/USD amounts on the customer PDF.
    final_label = meta.get("_final_price_label", "")
    final_value = meta.get("_final_price_value", "")

    if final_value:
        price_label_text = final_label or "TOTAL PACKAGE COST"
        amount_display = final_value
    else:
        # Price workflow not yet run — generate a reasonable placeholder
        nights_num = 0
        try:
            nights_num = int(re.search(r'(\d+)\s*[Nn]ight', meta.get("nights", "0")).group(1))
        except Exception:
            pass
        persons_num = 1
        try:
            persons_num = max(1, int(re.search(r'(\d+)\s*[Aa]dult', meta.get("persons", "1")).group(1)))
        except Exception:
            pass
        # ~INR 5,000 per night per person as placeholder
        placeholder_amt = max(15000, nights_num * persons_num * 5000)
        # Round to nearest 5000
        placeholder_amt = (placeholder_amt // 5000) * 5000
        amount_display = f"INR {placeholder_amt:,}/- (approx.)"
        price_label_text = "TOTAL PACKAGE COST"

    # ── Package price HTML (uses computed amount with fallback) ────────────────
    persons_rooms = "   ·   ".join(filter(None, [
        meta.get("persons", ""), meta.get("transport", "")]))
    price_html = f"""
        <div class="price-card">
          <div class="price-label">{_h(price_label_text.upper())}</div>
          <div class="price">{_h(str(amount_display))}</div>
          <div class="price-meta">{_h(persons_rooms)}</div>
        </div>"""


    # ── Day cards ─────────────────────────────────────────────────────────────
    days_html = ""
    all_days = meta.get("days", [])
    last_day_num = all_days[-1].get("day") if all_days else None
    for day in all_days:
        acts = "".join(f"<li>{_h(a)}</li>" for a in day.get("activities", []))
        # Suppress "Overnight in..." on the final/departure day
        is_last = day.get("day") == last_day_num
        overnight_city = day.get("overnight", "")
        overnight = (f'<div class="overnight">Overnight in {_h(overnight_city)}</div>'
                     if overnight_city and not is_last else "")
        days_html += f"""
        <div class="day-card">
          <div class="day-num">
            <div class="day-num-val">{_h(str(day.get('day','')))}</div>
            <div class="day-num-date">{_h(str(day.get('date','')).upper())}</div>
          </div>
          <div class="day-body">
            <div class="day-title">{_h(day.get('title',''))}</div>
            <ul>{acts}</ul>
            {overnight}
          </div>
        </div>"""

    # ── Hotels table ──────────────────────────────────────────────────────────
    hotels_html = ""
    if meta.get("hotels"):
        rows = "".join(
            f'<tr><td>{_h(h.get("city",""))}</td>'
            f'<td>{_h(h.get("hotel",""))}</td>'
            f'<td>{_h(h.get("dates",""))}</td></tr>'
            for h in meta["hotels"]
        )
        hotels_html = f"""
        <div class="section">
          <div class="section-title">Hotels Confirmed</div>
          <table class="hotel-table">
            <thead><tr>
              <th>City / Nights</th><th>Hotel</th><th>Dates</th>
            </tr></thead>
            <tbody>{rows}</tbody>
          </table>
        </div>"""

    # ── Highlights pills ──────────────────────────────────────────────────────
    highlights_html = ""
    if meta.get("highlights"):
        pills = "".join(f'<span class="hi-pill">{_h(h)}</span>'
                        for h in meta["highlights"])
        highlights_html = f"""
        <div class="section">
          <div class="section-title">Package Highlights</div>
          <div class="hi-pills">{pills}</div>
        </div>"""

    # ── Inclusions / Exclusions ───────────────────────────────────────────────
    inc = meta.get("inclusions", [])
    exc = meta.get("exclusions", [])
    ie_html = ""
    if inc or exc:
        inc_items = "".join(f"<li>{_h(i)}</li>" for i in inc)
        exc_items = "".join(f"<li>{_h(e)}</li>" for e in exc)
        ie_html = f"""
        <div class="ie-row">
          <div class="inc-card">
            <div class="inc-title"><span class="chk">&#10003;</span> Inclusions</div>
            <ul>{inc_items}</ul>
          </div>
          <div class="exc-card">
            <div class="exc-title"><span class="xmark">&#10005;</span> Exclusions</div>
            <ul>{exc_items}</ul>
          </div>
        </div>"""

    # ── Notes ─────────────────────────────────────────────────────────────────
    notes = meta.get("notes", [])
    notes_html = ""
    if notes:
        items = "".join(f"<li>{_h(n)}</li>" for n in notes)
        notes_html = f"""

        <div class="notes-card">
          <div class="notes-title">⚑ Important Notes</div>
          <ul>{items}</ul>
        </div>"""


    # ── DejaVu font paths ──────────────────────────────────────────────────────
    dv_dir = "/usr/share/fonts/truetype/dejavu"
    dv_regular = Path(f"{dv_dir}/DejaVuSans.ttf")
    dv_bold    = Path(f"{dv_dir}/DejaVuSans-Bold.ttf")

    font_face = ""
    if dv_regular.exists() and dv_bold.exists():
        r_b64 = base64.b64encode(dv_regular.read_bytes()).decode()
        b_b64 = base64.b64encode(dv_bold.read_bytes()).decode()
        font_face = f"""
        @font-face {{
            font-family: "DejaVuSans";
            font-weight: normal;
            src: url("data:font/truetype;base64,{r_b64}");
        }}
        @font-face {{
            font-family: "DejaVuSans";
            font-weight: bold;
            src: url("data:font/truetype;base64,{b_b64}");
        }}"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>
{font_face}

@page {{
    size: A4;
    margin: 14mm 16mm;
}}

* {{
    box-sizing: border-box;
    margin: 0;
    padding: 0;
}}

body {{
    font-family: "DejaVuSans", Arial, sans-serif;
    font-size: 11px;
    color: #333333;
    background: #FFFFFF;
}}

.page {{
    width: 100%;
    background: #FFFFFF;
}}

/* ── HERO ─────────────────────────────────────────────── */
.hero {{
    border-radius: 7px;
    overflow: hidden;
    height: 60mm !important;
    min-height: 60mm !important;
    max-height: 60mm !important;
    flex-shrink: 0;
    background: {hero_bg};
    background-size: cover !important;
    background-position: center center !important;
    background-repeat: no-repeat !important;
    display: flex;
    align-items: center;
    justify-content: center;
    text-align: center;
    margin-bottom: 10px;
    page-break-after: avoid;
    break-after: avoid;
}}
.hero-content {{
    width: 100%;
    padding: 20px 24px;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    text-align: center;
}}
.hero-title {{
    font-size: 25px;
    font-weight: bold;
    letter-spacing: 0.5px;
    color: #FFFFFF;
    text-transform: uppercase;
    line-height: 1.2;
    text-align: center;
    margin-left: auto;
    margin-right: auto;
}}
.hero-route {{
    font-size: 13px;
    color: #FFFFFF;
    margin-top: 6px;
    opacity: 0.9;
    text-align: center;
    margin-left: auto;
    margin-right: auto;
}}
.hero-dates {{
    font-size: 11px;
    color: #FFFFFF;
    margin-top: 5px;
    opacity: 0.85;
    text-align: center;
    margin-left: auto;
    margin-right: auto;
}}
.hero-logo {{
    display: block;
    margin: 0 auto 10px auto;
    height: 64px;
    width: auto;
    max-width: 320px;
    object-fit: contain;
    opacity: 0.95;
    filter: drop-shadow(0 1px 3px rgba(0,0,0,0.45));
}}

/* ── PILLS ────────────────────────────────────────────── */
.pills {{
    width: 100%;
    display: flex;
    flex-wrap: wrap;
    justify-content: center;
    align-items: center;
    gap: 5px;
    margin-bottom: 12px;
}}
.pill {{
    display: inline-flex;
    align-items: center;
    justify-content: center;
    padding: 6px 13px;
    border-radius: 999px;
    background: {theme['pill_bg']};
    border: 1px solid {theme['pill_border']};
    color: {theme['pill_text']};
    font-size: 10px;
    font-weight: bold;
    text-align: center;
}}


/* ── SECTION HEADINGS ─────────────────────────────────── */
.section {{
    margin-bottom: 12px;
}}
.section-title {{
    font-size: 17px;
    font-weight: bold;
    color: {theme['primary']};
    margin: 10px 0 7px 0;
    text-align: left;
}}

/* ── DAY CARDS ────────────────────────────────────────── */
.day-card {{
    display: flex;
    width: 100%;
    background: {theme['bg_light']};
    border: 1px solid {theme['border']};
    border-radius: 7px;
    overflow: hidden;
    margin-bottom: 7px;
    break-inside: avoid;
    page-break-inside: avoid;
}}
.day-num {{
    width: 52px;
    min-width: 52px;
    background: {theme['primary']};
    color: #FFFFFF;

    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    padding: 10px 4px;
    text-align: center;
}}
.day-num-val {{
    font-size: 20px;
    font-weight: bold;
    line-height: 1;
}}
.day-num-date {{
    font-size: 8px;
    font-weight: bold;
    margin-top: 3px;
    opacity: 0.85;
}}
.day-body {{
    padding: 10px 13px;
    flex: 1;
}}
.day-title {{
    font-size: 13px;
    font-weight: bold;
    color: {theme['primary']};
    margin-bottom: 5px;
}}
.day-body ul {{
    padding-left: 14px;
    margin: 0;
}}
.day-body li {{
    font-size: 10.5px;
    line-height: 1.5;
    margin-bottom: 2px;
    color: {theme['text']};
}}
.overnight {{
    font-size: 10px;
    font-weight: bold;
    color: {theme['primary']};
    margin-top: 6px;
}}

/* ── HOTEL TABLE ──────────────────────────────────────── */
.hotel-table {{
    width: 100%;
    border-collapse: separate;
    border-spacing: 0;
    overflow: hidden;
    border-radius: 7px;
    border: 1px solid {theme['border']};
}}
.hotel-table thead tr {{
    background: {theme['primary']};
}}
.hotel-table th {{
    padding: 8px 9px;
    font-size: 10px;
    font-weight: bold;
    color: #FFFFFF;
    text-align: left;
}}
.hotel-table td {{
    padding: 7px 9px;
    font-size: 10px;
    color: {theme['muted_text']};
    border-top: 1px solid {theme['border']};
}}
.hotel-table tr:nth-child(even) td {{
    background: {theme['hotel_row_alt']};
}}

/* ── PRICE CARD ───────────────────────────────────────── */
.price-card {{
    background: linear-gradient(135deg,{theme['primary']},{theme['primary_light']});
    color: #FFFFFF;

    border-radius: 7px;
    padding: 13px 20px;
    text-align: center;
    margin-top: 10px;
    break-inside: avoid;
    page-break-inside: avoid;
}}
.price-label {{
    font-size: 9px;
    letter-spacing: 1.2px;
    text-transform: uppercase;
    opacity: 0.85;
}}
.price {{
    font-size: 24px;
    font-weight: bold;
    margin: 5px 0 3px;
}}
.price-meta {{
    font-size: 9px;
    opacity: 0.8;
}}

/* ── HIGHLIGHTS ───────────────────────────────────────── */
.hi-pills {{
    display: flex;
    flex-wrap: wrap;
    gap: 5px;
}}
.hi-pill {{
    display: inline-block;
    padding: 6px 11px;
    border-radius: 999px;
    background: #FFF3DD;
    border: 1px solid #EAD8B4;
    color: #8A5A20;
    font-size: 9.5px;
    font-weight: bold;
}}

/* ── INCLUSIONS / EXCLUSIONS ──────────────────────────── */
.ie-row {{
    display: flex;
    gap: 8px;
    break-inside: avoid;
    page-break-inside: avoid;
}}
.inc-card {{
    flex: 1;
    background: #EFFAF1;
    border: 1px solid #CDE8D1;
    border-radius: 7px;
    padding: 10px 12px;
}}
.exc-card {{
    flex: 1;
    background: #FFF1F1;
    border: 1px solid #F0D1D1;
    border-radius: 7px;
    padding: 10px 12px;
}}
.inc-title {{
    color: #3B7D44;
    font-size: 12px;
    font-weight: bold;
    margin-bottom: 6px;
}}
.exc-title {{
    color: #A52A2A;
    font-size: 12px;
    font-weight: bold;
    margin-bottom: 6px;
}}
.chk {{
    display: inline-block;
    color: #3B7D44;
    font-weight: bold;
    font-size: 13px;
    line-height: 1;
    margin-right: 3px;
}}
.xmark {{
    display: inline-block;
    color: #A52A2A;
    font-weight: bold;
    font-size: 13px;
    line-height: 1;
    margin-right: 3px;
}}
.inc-card ul, .exc-card ul {{
    padding-left: 13px;
    margin: 0;
}}
.inc-card li, .exc-card li {{
    font-size: 10px;
    line-height: 1.55;
    margin-bottom: 2px;
}}

/* ── NOTES ────────────────────────────────────────────── */
.notes-card {{
    background: #FFF8E7;
    border: 1px solid #EEDFB5;
    border-radius: 7px;
    padding: 10px 13px;
    margin-top: 9px;
    break-inside: avoid;
    page-break-inside: avoid;
}}
.notes-title {{
    color: #8A651C;
    font-size: 12px;
    font-weight: bold;
    margin-bottom: 6px;
}}
.notes-card ul {{
    padding-left: 13px;
    margin: 0;
}}
.notes-card li {{
    font-size: 10px;
    line-height: 1.55;
    margin-bottom: 2px;
    color: #555533;
}}

/* ── CONTACT FOOTER ───────────────────────────────────── */
.contact-footer {{
    background: linear-gradient(135deg,{theme['primary']},{theme['primary_light']});
    color: #FFFFFF;

    border-radius: 7px;
    padding: 11px 16px;
    text-align: center;
    margin-top: 12px;
    break-inside: avoid;
    page-break-inside: avoid;
}}
.agent-name {{
    font-size: 14px;
    font-weight: bold;
    margin-bottom: 3px;
}}
.contact-details {{
    font-size: 9px;
    font-weight: bold;
    opacity: 0.88;
}}
</style>
</head>
<body>
<div class="page">

  <!-- HERO -->
  <div class="hero">
    <div class="hero-content">
      {logo_tag}
      <div class="hero-title">{pkg_name}</div>
      {f'<div class="hero-route">{route_txt}</div>' if route_txt else ''}
      {f'<div class="hero-dates">{hero_dates}</div>' if hero_dates else ''}
    </div>
  </div>

  <!-- PILLS -->
  <div class="pills">{pills_html}</div>

  <!-- DAY-WISE ITINERARY -->
  <div class="section">
    <div class="section-title">Day-wise Itinerary</div>
    {days_html}
  </div>

  {hotels_html}
  {price_html}
  {highlights_html}

  <!-- INCLUSIONS / EXCLUSIONS -->
  {ie_html}

  {notes_html}

  <!-- CONTACT FOOTER -->
  <div class="contact-footer">
    <div class="agent-name">{_h(COMPANY_NAME)}</div>
    <div class="contact-details">{_h(COMPANY_PHONE)} &nbsp;·&nbsp; {_h(COMPANY_WEB)} &nbsp;·&nbsp; {_h(COMPANY_EMAIL)}</div>
  </div>

</div>
</body>
</html>"""

    buf = io.BytesIO()
    WH(string=html).write_pdf(buf)
    return buf.getvalue()


def generate_pdf_editorial(title: str, meta: dict, theme: dict | None = None) -> bytes:
    """
    Render the "Modern Editorial" layout — a clean, document-style itinerary
    inspired visually by the reference Georgia_Tbilisi_Grand_Explorer PDF:
    compact trip-summary strip, a plain package-title block (no photo hero),
    large typographic day numbers with a thin accent rule, a simple hotel
    table, and restrained accent-color usage throughout. This is a
    completely different visual system from generate_pdf()'s "classic"
    layout, but it is driven by the EXACT SAME `meta` dict — no itinerary
    content, hotel data, pricing, inclusions/exclusions, or notes differ
    between layouts; only the presentation differs.

    Uses local DejaVu Sans fonts — no internet required. No hero photo is
    used in this layout (by design, matching the reference PDF).
    """
    import base64
    from weasyprint import HTML as WH

    theme = theme or PDF_THEMES["Modern Editorial"]

    pkg_name   = _h(meta.get("package_name", meta.get("destination", title)).upper())
    route_txt  = _h(meta.get("route", ""))
    dates_txt  = _h(meta.get("dates", ""))
    nights_txt = _h(_format_nights_days(meta.get("nights", ""), meta.get("days")))

    # ── Compact trip-summary strip ──────────────────────────────────────────
    summary_items = [
        meta.get("dates", ""), _format_nights_days(meta.get("nights", ""), meta.get("days")),
        meta.get("persons", ""), meta.get("transport", ""),
    ]
    summary_html = "".join(
        f'<span class="sum-item">{_h(s)}</span>' for s in summary_items if s
    )

    # ── Final customer-facing price (same logic as classic layout) ─────────
    final_label = meta.get("_final_price_label", "")
    final_value = meta.get("_final_price_value", "")
    if final_value:
        price_label_text = final_label or "TOTAL PACKAGE COST"
        amount_display = final_value
    else:
        nights_num = 0
        try:
            nights_num = int(re.search(r'(\d+)\s*[Nn]ight', meta.get("nights", "0")).group(1))
        except Exception:
            pass
        persons_num = 1
        try:
            persons_num = max(1, int(re.search(r'(\d+)\s*[Aa]dult', meta.get("persons", "1")).group(1)))
        except Exception:
            pass
        placeholder_amt = max(15000, nights_num * persons_num * 5000)
        placeholder_amt = (placeholder_amt // 5000) * 5000
        amount_display = f"INR {placeholder_amt:,}/- (approx.)"
        price_label_text = "TOTAL PACKAGE COST"

    # ── Day blocks — large typographic day number + thin accent rule ───────
    days_html = ""
    all_days = meta.get("days", [])
    last_day_num = all_days[-1].get("day") if all_days else None
    for day in all_days:
        acts = "".join(f"<li>{_h(a)}</li>" for a in day.get("activities", []))
        is_last = day.get("day") == last_day_num
        overnight_city = day.get("overnight", "")
        overnight = (f'<div class="ed-overnight">Overnight in {_h(overnight_city)}</div>'
                     if overnight_city and not is_last else "")
        day_date = _h(str(day.get('date', '')))
        days_html += f"""
        <div class="ed-day">
          <div class="ed-day-num-col">
            <div class="ed-day-num">{_h(str(day.get('day', '')).zfill(2))}</div>
            <div class="ed-day-label">DAY {_h(str(day.get('day', '')))}</div>
          </div>
          <div class="ed-day-body">
            <div class="ed-day-title">{_h(day.get('title', ''))}{f' <span class="ed-day-date">— {day_date}</span>' if day_date else ''}</div>
            <ul>{acts}</ul>
            {overnight}
          </div>
        </div>"""

    # ── Hotel Details table ──────────────────────────────────────────────────
    hotels_html = ""
    if meta.get("hotels"):
        rows = "".join(
            f'<tr><td>{_h(h.get("city",""))}</td>'
            f'<td>{_h(h.get("hotel",""))}</td>'
            f'<td>{_h(h.get("dates",""))}</td></tr>'
            for h in meta["hotels"]
        )
        hotels_html = f"""
        <div class="ed-section">
          <div class="ed-section-title">Hotel Details</div>
          <table class="ed-table">
            <thead><tr><th>City / Nights</th><th>Hotel Options</th><th>Dates</th></tr></thead>
            <tbody>{rows}</tbody>
          </table>
        </div>"""

    # ── Highlights ────────────────────────────────────────────────────────
    highlights_html = ""
    if meta.get("highlights"):
        items = "".join(f"<li>{_h(h)}</li>" for h in meta["highlights"])
        highlights_html = f"""
        <div class="ed-section">
          <div class="ed-section-title">Package Highlights</div>
          <ul class="ed-highlights">{items}</ul>
        </div>"""

    # ── Inclusions / Exclusions ──────────────────────────────────────────────
    inc = meta.get("inclusions", [])
    exc = meta.get("exclusions", [])
    ie_html = ""
    if inc or exc:
        inc_items = "".join(f"<li>{_h(i)}</li>" for i in inc)
        exc_items = "".join(f"<li>{_h(e)}</li>" for e in exc)
        ie_html = f"""
        <div class="ed-ie-row">
          <div class="ed-inc">
            <div class="ed-ie-title">&#10003; Inclusions</div>
            <ul>{inc_items}</ul>
          </div>
          <div class="ed-exc">
            <div class="ed-ie-title">&#10005; Exclusions</div>
            <ul>{exc_items}</ul>
          </div>
        </div>"""

    # ── Important Notes ──────────────────────────────────────────────────────
    notes = meta.get("notes", [])
    notes_html = ""
    if notes:
        items = "".join(f"<li>{_h(n)}</li>" for n in notes)
        notes_html = f"""
        <div class="ed-notes">
          <div class="ed-notes-title">Important Notes</div>
          <ul>{items}</ul>
        </div>"""

    # ── DejaVu font paths ──────────────────────────────────────────────────
    dv_dir = "/usr/share/fonts/truetype/dejavu"
    dv_regular = Path(f"{dv_dir}/DejaVuSans.ttf")
    dv_bold    = Path(f"{dv_dir}/DejaVuSans-Bold.ttf")
    font_face = ""
    if dv_regular.exists() and dv_bold.exists():
        r_b64 = base64.b64encode(dv_regular.read_bytes()).decode()
        b_b64 = base64.b64encode(dv_bold.read_bytes()).decode()
        font_face = f"""
        @font-face {{ font-family: "DejaVuSans"; font-weight: normal;
            src: url("data:font/truetype;base64,{r_b64}"); }}
        @font-face {{ font-family: "DejaVuSans"; font-weight: bold;
            src: url("data:font/truetype;base64,{b_b64}"); }}"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>
{font_face}

@page {{ size: A4; margin: 16mm 18mm; }}
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{
    font-family: "DejaVuSans", Arial, sans-serif;
    font-size: 10.5px;
    color: {theme['text']};
    background: #FFFFFF;
}}

/* ── PACKAGE TITLE BLOCK ─────────────────────────────────── */
.ed-title {{
    font-size: 22px;
    font-weight: bold;
    letter-spacing: 0.3px;
    color: {theme['primary']};
    text-transform: uppercase;
    margin-bottom: 3px;
}}
.ed-route {{
    font-size: 11px;
    color: {theme['muted_text']};
    margin-bottom: 3px;
}}
.ed-dates-line {{
    font-size: 10.5px;
    color: {theme['accent']};
    font-weight: bold;
    margin-bottom: 10px;
}}
.ed-rule {{
    height: 2px;
    background: {theme['accent']};
    width: 46px;
    margin-bottom: 12px;
}}

/* ── SUMMARY STRIP ────────────────────────────────────────── */
.ed-summary {{
    display: flex;
    flex-wrap: wrap;
    gap: 0;
    border-top: 1px solid {theme['border']};
    border-bottom: 1px solid {theme['border']};
    padding: 8px 0;
    margin-bottom: 14px;
}}
.sum-item {{
    font-size: 9.5px;
    font-weight: bold;
    color: {theme['text']};
    padding: 0 12px;
    border-right: 1px solid {theme['border']};
}}
.sum-item:last-child {{ border-right: none; }}

/* ── PRICE ────────────────────────────────────────────────── */
.ed-price {{
    text-align: right;
    margin-bottom: 16px;
}}
.ed-price-label {{
    font-size: 8.5px;
    letter-spacing: 1px;
    text-transform: uppercase;
    color: {theme['muted_text']};
}}
.ed-price-val {{
    font-size: 20px;
    font-weight: bold;
    color: {theme['accent']};
}}

/* ── SECTION HEADINGS ─────────────────────────────────────── */
.ed-section {{ margin-bottom: 16px; }}
.ed-section-title {{
    font-size: 14px;
    font-weight: bold;
    color: {theme['primary']};
    text-transform: uppercase;
    letter-spacing: 0.5px;
    padding-bottom: 5px;
    border-bottom: 1.5px solid {theme['accent']};
    margin-bottom: 9px;
}}

/* ── HOTEL TABLE ──────────────────────────────────────────── */
.ed-table {{
    width: 100%;
    border-collapse: collapse;
}}
.ed-table th {{
    text-align: left;
    font-size: 9.5px;
    font-weight: bold;
    color: #FFFFFF;
    background: {theme['primary']};
    padding: 7px 9px;
}}
.ed-table td {{
    font-size: 9.5px;
    color: {theme['text']};
    padding: 6px 9px;
    border-bottom: 1px solid {theme['border']};
}}
.ed-table tr:nth-child(even) td {{ background: {theme['hotel_row_alt']}; }}

/* ── DAY-WISE ITINERARY ───────────────────────────────────── */
.ed-day {{
    display: flex;
    gap: 14px;
    padding: 12px 0;
    border-bottom: 1px solid {theme['border']};
    break-inside: avoid;
    page-break-inside: avoid;
}}
.ed-day-num-col {{
    width: 54px;
    min-width: 54px;
    text-align: center;
}}
.ed-day-num {{
    font-size: 30px;
    font-weight: bold;
    line-height: 1;
    color: {theme['accent']};
}}
.ed-day-label {{
    font-size: 8px;
    font-weight: bold;
    letter-spacing: 1px;
    color: {theme['muted_text']};
    margin-top: 2px;
}}
.ed-day-body {{ flex: 1; padding-top: 2px; }}
.ed-day-title {{
    font-size: 13px;
    font-weight: bold;
    color: {theme['primary']};
    margin-bottom: 6px;
}}
.ed-day-date {{
    font-size: 10px;
    font-weight: normal;
    color: {theme['muted_text']};
}}
.ed-day-body ul {{ padding-left: 14px; margin: 0; }}
.ed-day-body li {{
    font-size: 10px;
    line-height: 1.55;
    margin-bottom: 3px;
    color: {theme['text']};
}}
.ed-overnight {{
    font-size: 9.5px;
    font-weight: bold;
    color: {theme['accent']};
    margin-top: 6px;
}}

/* ── HIGHLIGHTS ───────────────────────────────────────────── */
.ed-highlights {{ padding-left: 16px; }}
.ed-highlights li {{
    font-size: 10px;
    line-height: 1.6;
    color: {theme['text']};
    margin-bottom: 3px;
}}

/* ── INCLUSIONS / EXCLUSIONS ──────────────────────────────── */
.ed-ie-row {{
    display: flex;
    gap: 16px;
    break-inside: avoid;
    page-break-inside: avoid;
}}
.ed-inc, .ed-exc {{ flex: 1; }}
.ed-ie-title {{
    font-size: 11.5px;
    font-weight: bold;
    color: {theme['primary']};
    margin-bottom: 6px;
}}
.ed-inc ul, .ed-exc ul {{ padding-left: 14px; margin: 0; }}
.ed-inc li, .ed-exc li {{
    font-size: 9.5px;
    line-height: 1.6;
    color: {theme['text']};
    margin-bottom: 3px;
}}

/* ── IMPORTANT NOTES ──────────────────────────────────────── */
.ed-notes {{
    margin-top: 6px;
    padding-top: 10px;
    border-top: 1px solid {theme['border']};
    break-inside: avoid;
    page-break-inside: avoid;
}}
.ed-notes-title {{
    font-size: 11.5px;
    font-weight: bold;
    color: {theme['primary']};
    margin-bottom: 6px;
}}
.ed-notes ul {{ padding-left: 14px; margin: 0; }}
.ed-notes li {{
    font-size: 9px;
    line-height: 1.6;
    color: {theme['muted_text']};
    margin-bottom: 3px;
}}

/* ── FOOTER ───────────────────────────────────────────────── */
.ed-footer {{
    margin-top: 16px;
    padding-top: 10px;
    border-top: 2px solid {theme['primary']};
    text-align: center;
}}
.ed-agent-name {{
    font-size: 12px;
    font-weight: bold;
    color: {theme['primary']};
    margin-bottom: 3px;
}}
.ed-contact {{
    font-size: 9px;
    color: {theme['muted_text']};
}}
</style>
</head>
<body>

  <div class="ed-title">{pkg_name}</div>
  {f'<div class="ed-route">{route_txt}</div>' if route_txt else ''}
  {f'<div class="ed-dates-line">{dates_txt or nights_txt}{" | " + nights_txt if dates_txt and nights_txt else ""}</div>' if (dates_txt or nights_txt) else ''}
  <div class="ed-rule"></div>

  <div class="ed-summary">{summary_html}</div>

  <div class="ed-price">
    <div class="ed-price-label">{_h(price_label_text.upper())}</div>
    <div class="ed-price-val">{_h(str(amount_display))}</div>
  </div>

  {hotels_html}

  <div class="ed-section">
    <div class="ed-section-title">Day-wise Itinerary</div>
    {days_html}
  </div>

  {highlights_html}
  {ie_html}
  {notes_html}

  <div class="ed-footer">
    <div class="ed-agent-name">{_h(COMPANY_NAME)}</div>
    <div class="ed-contact">{_h(COMPANY_PHONE)} &nbsp;·&nbsp; {_h(COMPANY_WEB)} &nbsp;·&nbsp; {_h(COMPANY_EMAIL)}</div>
  </div>

</body>
</html>"""

    buf = io.BytesIO()
    WH(string=html).write_pdf(buf)
    return buf.getvalue()


def generate_pdf_themed(title: str, meta: dict, theme_name: str,
                        header_img_path: Path | None = None) -> bytes:
    """
    Single entry point PDF generation dispatches to based on the selected
    theme's declared layout — "classic" (generate_pdf, gradient hero photo)
    or "editorial" (generate_pdf_editorial, document-style, no photo).
    Callers should use this function rather than calling generate_pdf() or
    generate_pdf_editorial() directly, so that adding a new theme with a
    new layout only requires updating this dispatcher.
    """
    theme = PDF_THEMES.get(theme_name, PDF_THEMES[DEFAULT_PDF_THEME])
    if theme.get("layout") == "editorial":
        return generate_pdf_editorial(title, meta, theme=theme)
    return generate_pdf(title, "", meta, header_img_path=header_img_path, theme=theme)


# ── LLM-generated HTML → PDF ─────────────────────────────────────────────────


def generate_html_pdf(meta: dict, api_key: str, model: str,
                      header_img_path: Path | None = None) -> bytes:
    """Ask Gemini to design a modern HTML itinerary, then render to PDF via WeasyPrint."""
    import base64, requests
    from weasyprint import HTML as WH, CSS

    # Embed header image as base64 if available
    img_tag = '<div class="hero-fallback"></div>'
    if header_img_path and header_img_path.exists():
        img_b64 = base64.b64encode(header_img_path.read_bytes()).decode()
        img_tag = f'<img class="hero-img" src="data:image/jpeg;base64,{img_b64}" />'

    # Serialize meta for Gemini
    meta_json = json.dumps(meta, indent=2)

    system = (
        "You are a world-class HTML/CSS designer specialising in travel brochures. "
        "Generate a complete, self-contained HTML page for a travel itinerary PDF. "
        "Requirements:\n"
        "- Modern, professional design inspired by luxury travel brochures\n"
        "- Use CSS variables, gradients, and subtle shadows\n"
        "- Color palette: deep purple/navy (#2A1830) for accents, maroon (#8A1A1A) for headings, "
        "warm white (#FAFAF8) background, amber (#8A6300) for highlights\n"
        "- Page size A4, all content self-contained (no external URLs except fonts)\n"
        "- First section: full-width hero header with the destination image tag provided\n"
        "- Overlay destination name, route cities, and trip details on the hero image\n"
        "- Day cards: numbered boxes on the left with dark purple background, activities on the right\n"
        "- Hotels table, Package highlights grid, Inclusions/Exclusions two-column layout\n"
        "- Footer with company name, phone, web, email\n"
        "- Use Google Fonts (Playfair Display for headings, Inter for body) via @import in <style>\n"
        "- Include the hero image placeholder where indicated: {IMG_TAG}\n"
        "- Print-friendly: @media print rules, page breaks where appropriate\n"
        "- Output ONLY the complete HTML. No explanation. No markdown fences.\n\n"
        f"Company: {COMPANY_NAME} | {COMPANY_PHONE} | {COMPANY_WEB} | {COMPANY_EMAIL}\n\n"
        f"Itinerary data:\n{meta_json}\n\n"
        f"Replace {{IMG_TAG}} with: {img_tag}"
    )

    _p2 = os.environ.get("HTTPS_PROXY") or os.environ.get("HTTP_PROXY") or ""
    proxies = {"https": _p2, "http": _p2} if _p2 else {}
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    payload = {
        "system_instruction": {"parts": [{"text": system}]},
        "contents": [{"role": "user", "parts": [{"text": "Generate the travel itinerary HTML now."}]}],
        "generationConfig": {"temperature": 0.4, "maxOutputTokens": 8192}
    }
    resp = requests.post(url, json=payload, headers={"Content-Type": "application/json"},
                         params={"key": api_key}, proxies=proxies, timeout=120)
    resp.raise_for_status()
    html_text = resp.json()["candidates"][0]["content"]["parts"][0]["text"]

    # Strip markdown fences if Gemini wrapped it
    html_text = re.sub(r'^```[a-z]*\n?', '', html_text.strip())
    html_text = re.sub(r'\n?```$', '', html_text.strip())

    # Render HTML → PDF
    buf = io.BytesIO()
    WH(string=html_text).write_pdf(buf)
    return buf.getvalue()


# ── Document loaders ─────────────────────────────────────────────────────────

def load_pdf_text(path: Path) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)

def load_docx_text(path: Path) -> str:
    """
    Extract text from a DOCX in TRUE READING ORDER, interleaving paragraphs
    and tables exactly as they appear in the document body.

    This is critical for multi-package documents (e.g. several Georgia
    itinerary variants in one file) where each package has its own pricing
    table immediately following its package-name heading. Processing all
    paragraphs first and all tables afterward (the previous approach)
    completely disconnects each price table from its package heading,
    causing the LLM/retrieval to pull the wrong package's rate table.

    Each table row is also prefixed with the most recent package/section
    heading seen so far, so retrieved chunks always carry enough context
    to identify which package the rates/hotels belong to — even if the
    heading itself ends up in a different chunk after text splitting.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    doc = Document(path)
    body = doc.element.body

    parts = []
    current_heading = ""
    # Headings are short, non-tabular lines that look like a package/section
    # title — e.g. "6 Nights - 4 Nights Tbilisi + 2 Nights Batumi".
    # IMPORTANT: generic lines like "Duration: 7 Days and 6 Nights" or
    # "Validity: April 1 to October 31, 2026" also mention Nights/Days but
    # are NOT unique package identifiers — they repeat verbatim across
    # multiple different packages in this document. If these overwrite
    # current_heading, every subsequent price table gets tagged with the
    # wrong (non-unique) package label, causing rate mix-ups between
    # packages. Only the actual package-title line (which names the
    # cities/route, not just a generic duration/validity statement) should
    # update current_heading.
    heading_re = re.compile(r'\bNights?\b|\bDays?\b', re.I)
    day_re = re.compile(r'^\s*Day\s*\d+\s*:', re.I)
    non_title_re = re.compile(
        r'^\s*(Duration|Validity|Cities|Rates|Accommodation|Supplement|'
        r'No\.?\s*of\s*Pax|Destination)\s*[:\-]', re.I
    )

    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = Paragraph(child, doc)
            text = p.text.strip()
            if not text:
                continue
            parts.append(text)
            # Update current package heading — prefer lines like
            # "6 Nights - 4 Nights Tbilisi + 2 Nights Batumi" over
            # per-day headings like "Day 1: Arrival" or generic
            # "Duration: ..." / "Validity: ..." lines that repeat
            # identically across many different packages.
            if (heading_re.search(text) and not day_re.match(text)
                    and not non_title_re.match(text) and len(text) < 120):
                current_heading = text

        elif child.tag == qn('w:tbl'):
            table = Table(child, doc)
            if not table.rows:
                continue
            headers = [c.text.strip() for c in table.rows[0].cells]
            for row in table.rows[1:]:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                if not any(cells):
                    continue
                if headers:
                    row_text = " | ".join(
                        f"{h}: {v}" for h, v in zip(headers, cells) if v
                    )
                else:
                    row_text = " | ".join(c for c in cells if c)
                if current_heading:
                    row_text = f"[Package: {current_heading}] {row_text}"
                parts.append(row_text)

    return "\n".join(parts)


def extract_all_price_tables(docs_dir: Path) -> str:
    """
    Extract EVERY price/rate table from every DOCX in docs_dir, tagged with
    its package heading, and return them concatenated as a single block.

    WHY THIS EXISTS: with 6+ near-identical pricing tables in one document
    (same headers "No of Pax | 5 Star | 4 Star... | 3 Star", same row
    labels "2 Pax / 4 Pax / 6 Pax"), semantic vector similarity search
    cannot reliably distinguish which table belongs to which package —
    it may retrieve only PART of a table, or rows from the wrong package,
    since the embeddings for two different packages' price rows are nearly
    identical. This caused wrong USD rates to be quoted (e.g. picking the
    "4 Nights Tbilisi" table's numbers for the "6 Nights - 4 Nights Tbilisi
    + 2 Nights Batumi" package).

    FIX: bypass retrieval entirely for pricing. Extract the complete,
    authoritative set of price tables directly from the source documents
    (already tagged with unique package headings by load_docx_text's
    reading-order walk) and inject the FULL, UNTRUNCATED block into every
    LLM call's context. This guarantees the model always sees the
    complete, correctly-labeled table for every package, rather than a
    similarity-searched fragment that might belong to a different package.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    heading_re = re.compile(r'\bNights?\b|\bDays?\b', re.I)
    day_re = re.compile(r'^\s*Day\s*\d+\s*:', re.I)
    non_title_re = re.compile(
        r'^\s*(Duration|Validity|Cities|Rates|Accommodation|Supplement|'
        r'No\.?\s*of\s*Pax|Destination)\s*[:\-]', re.I
    )
    # Only tables whose header row mentions Pax/Star/Cost are pricing tables
    price_header_re = re.compile(r'Pax|Star|Cost|Rate|Price|USD|INR', re.I)

    blocks = []
    for f in sorted(Path(docs_dir).glob("*.docx")):
        try:
            doc = Document(f)
        except Exception:
            continue
        body = doc.element.body
        current_heading = ""
        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                p = Paragraph(child, doc)
                text = p.text.strip()
                if not text:
                    continue
                if (heading_re.search(text) and not day_re.match(text)
                        and not non_title_re.match(text) and len(text) < 120):
                    current_heading = text
            elif child.tag == qn('w:tbl'):
                table = Table(child, doc)
                if not table.rows:
                    continue
                headers = [c.text.strip() for c in table.rows[0].cells]
                header_line = " ".join(headers)
                if not price_header_re.search(header_line):
                    continue  # skip non-pricing tables (e.g. hotel-name tables)
                rows_text = []
                for row in table.rows[1:]:
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    if not any(cells):
                        continue
                    row_text = " | ".join(
                        f"{h}: {v}" for h, v in zip(headers, cells) if v
                    )
                    rows_text.append(row_text)
                if rows_text and current_heading:
                    blocks.append(
                        f"### PRICE TABLE — Package: {current_heading} (source: {f.name})\n"
                        + "\n".join(rows_text)
                    )
    return "\n\n".join(blocks)


def extract_all_inclusion_exclusion_tables(docs_dir: Path) -> str:
    """
    Extract EVERY Inclusions/Exclusions table from every DOCX in docs_dir,
    tagged with its package heading, and return them concatenated as a
    single authoritative block.

    WHY THIS EXISTS: exactly the same problem as extract_all_price_tables —
    each package (e.g. "5 Nights - 4 Nights Tbilisi + 1 Night Gudauri" vs
    "6 Nights - 4 Nights Tbilisi + 2 Nights Batumi") has its OWN two-column
    "Inclusions | Exclusions" table with nights/city-specific line items
    (e.g. "3 Nights' accommodation in Tbilisi", "1 Night accommodation in
    Gudauri"), but the wording overlaps heavily across packages ("Daily
    Breakfast", "Round Airport Transfers", "Tips for guide and driver" appear
    almost everywhere). Semantic similarity search cannot reliably tell which
    package's table a retrieved chunk belongs to, so it can quote the wrong
    package's night-by-night accommodation breakdown.

    FIX: bypass retrieval for inclusions/exclusions entirely. Extract the
    complete, correctly-labeled table for every package directly from the
    source documents (matching each table to the most recent package
    heading, exactly like price tables) and inject the full block into every
    LLM call's context.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    heading_re = re.compile(r'\bNights?\b|\bDays?\b', re.I)
    day_re = re.compile(r'^\s*Day\s*\d+\s*:', re.I)
    non_title_re = re.compile(
        r'^\s*(Duration|Validity|Cities|Rates|Accommodation|Supplement|'
        r'No\.?\s*of\s*Pax|Destination)\s*[:\-]', re.I
    )

    # Internal/agent-facing line items that must NEVER appear in the
    # customer-facing PDF's Exclusions list. These describe agent-to-vendor
    # banking/remittance charges, not something the traveller is exposed to.
    excluded_line_patterns = [
        re.compile(r'transactional\s+charges.*remitting.*recovered\s+from\s+the\s+agent', re.I),
    ]

    def _is_suppressed(line: str) -> bool:
        return any(p.search(line) for p in excluded_line_patterns)

    blocks = []
    for f in sorted(Path(docs_dir).glob("*.docx")):
        try:
            doc = Document(f)
        except Exception:
            continue
        body = doc.element.body
        current_heading = ""
        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                p = Paragraph(child, doc)
                text = p.text.strip()
                if not text:
                    continue
                if (heading_re.search(text) and not day_re.match(text)
                        and not non_title_re.match(text) and len(text) < 120):
                    current_heading = text
            elif child.tag == qn('w:tbl'):
                table = Table(child, doc)
                if not table.rows:
                    continue
                headers = [c.text.strip() for c in table.rows[0].cells]
                header_line = " ".join(headers).lower()
                # Only the 2-column "Inclusions | Exclusions" table
                if "inclusion" not in header_line or "exclusion" not in header_line:
                    continue
                inc_items, exc_items = [], []
                for row in table.rows[1:]:
                    cells = row.cells
                    if len(cells) >= 2:
                        inc_items += [
                            ln.strip(" \u00a0\u2022-")
                            for ln in cells[0].text.replace("\u00a0", " ").split("\n")
                            if ln.strip(" \u00a0\u2022-") and not _is_suppressed(ln)
                        ]
                        exc_items += [
                            ln.strip(" \u00a0\u2022-")
                            for ln in cells[1].text.replace("\u00a0", " ").split("\n")
                            if ln.strip(" \u00a0\u2022-") and not _is_suppressed(ln)
                        ]
                if (inc_items or exc_items) and current_heading:

                    block = (
                        f"### INCLUSIONS & EXCLUSIONS — Package: {current_heading} "
                        f"(source: {f.name})\n"
                        "Inclusions:\n" + "\n".join(f"- {i}" for i in inc_items) + "\n"
                        "Exclusions:\n" + "\n".join(f"- {e}" for e in exc_items)
                    )
                    blocks.append(block)
    return "\n\n".join(blocks)


def extract_all_hotel_tables(docs_dir: Path) -> str:
    """
    Extract EVERY "Accommodation Details" (hotel options per city/category)
    table from every DOCX in docs_dir, tagged with its package heading, and
    return them concatenated as a single authoritative block.

    WHY THIS EXISTS: same rationale as price tables and inclusions/exclusions
    — every package variant has its OWN accommodation table (same header row
    "Destination | 5 Star | 4 Star – (D) | ... | 3 Star", differing only in
    which cities are listed and which hotel names appear in the 3-Star / 4-
    Star / 5-Star cells for each city). Semantic retrieval cannot reliably
    keep these straight across 6 nearly-identical tables, so bypass retrieval
    entirely and inject the correctly-labeled table for every package.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    heading_re = re.compile(r'\bNights?\b|\bDays?\b', re.I)
    day_re = re.compile(r'^\s*Day\s*\d+\s*:', re.I)
    non_title_re = re.compile(
        r'^\s*(Duration|Validity|Cities|Rates|Accommodation|Supplement|'
        r'No\.?\s*of\s*Pax|Destination)\s*[:\-]', re.I
    )
    # Accommodation table's header row starts with "Destination" and lists
    # star categories — distinct from the pricing table (which starts with
    # "No of Pax") even though both use similar category-name columns.
    hotel_header_re = re.compile(r'Destination', re.I)

    blocks = []
    for f in sorted(Path(docs_dir).glob("*.docx")):
        try:
            doc = Document(f)
        except Exception:
            continue
        body = doc.element.body
        current_heading = ""
        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                p = Paragraph(child, doc)
                text = p.text.strip()
                if not text:
                    continue
                if (heading_re.search(text) and not day_re.match(text)
                        and not non_title_re.match(text) and len(text) < 120):
                    current_heading = text
            elif child.tag == qn('w:tbl'):
                table = Table(child, doc)
                if not table.rows:
                    continue
                headers = [c.text.strip() for c in table.rows[0].cells]
                if not headers or not hotel_header_re.search(headers[0]):
                    continue  # not the accommodation table
                rows_text = []
                for row in table.rows[1:]:
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    if not any(cells):
                        continue
                    row_text = " | ".join(
                        f"{h}: {v}" for h, v in zip(headers, cells) if v
                    )
                    rows_text.append(row_text)
                if rows_text and current_heading:
                    blocks.append(
                        f"### HOTEL OPTIONS — Package: {current_heading} (source: {f.name})\n"
                        + "\n".join(rows_text)
                    )
    return "\n\n".join(blocks)


def extract_all_important_notes(docs_dir: Path) -> str:
    """
    Extract EVERY "Important Notes:" free-text block from every DOCX in
    docs_dir, tagged with its package heading, and return them concatenated
    as a single authoritative block.

    WHY THIS EXISTS: "Important Notes" is free-flowing paragraph text (not a
    table), placed immediately after each package's Inclusions/Exclusions
    table. Because it is plain text rather than a distinctly-labeled table,
    the previous approach relied entirely on semantic chunk retrieval to
    surface it — which is unreliable across 6 near-identical packages and
    frequently caused these notes to be dropped from the generated PDF
    entirely. Bypass retrieval: walk the document in reading order, capture
    every paragraph following an "Important Notes" heading up to the next
    package heading (or an "INDEX" marker, or a table), and tag it with the
    current package.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    # NOTE: unlike other extractors in this file, we CANNOT use the loose
    # `heading_re = r'\bNights?\b|\bDays?\b'` pattern here. That pattern
    # matches "Nights"/"Days" ANYWHERE in the text, including as an ordinary
    # word inside an Important Notes sentence — e.g. "Separate guide service
    # if required per day 60 $..." contains "day" and would be misdetected
    # as a new package heading, causing an immediate/empty flush of the
    # notes before any lines are captured (this was the root cause of
    # "Important Notes" silently disappearing for every package). Instead,
    # require the line to actually START with a number followed by
    # "Night(s)" — the real pattern every package title uses in this
    # document (e.g. "4 Nights Tbilisi", "6 Nights - 4 Nights Tbilisi + 2
    # Nights Batumi") — which free-text note sentences never do.
    package_heading_re = re.compile(r'^\s*\d+\s*Nights?\b', re.I)
    notes_heading_re = re.compile(r'^\s*Important\s*Notes\s*:?\s*$', re.I)
    stop_re = re.compile(r'^\s*INDEX\s*$', re.I)


    blocks = []
    for f in sorted(Path(docs_dir).glob("*.docx")):
        try:
            doc = Document(f)
        except Exception:
            continue
        body = doc.element.body
        current_heading = ""
        in_notes = False
        note_lines: list[str] = []

        def _flush():
            nonlocal note_lines, current_heading
            if note_lines and current_heading:
                blocks.append(
                    f"### IMPORTANT NOTES — Package: {current_heading} (source: {f.name})\n"
                    + "\n".join(f"- {n}" for n in note_lines)
                )
            note_lines = []

        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                p = Paragraph(child, doc)
                text = p.text.strip()
                if not text:
                    continue
                if notes_heading_re.match(text):
                    in_notes = True
                    continue
                if package_heading_re.match(text) and len(text) < 120:
                    # A new package heading ends the current notes block
                    if in_notes:
                        _flush()
                        in_notes = False
                    current_heading = text
                    continue
                if stop_re.match(text):
                    if in_notes:
                        _flush()
                        in_notes = False
                    continue
                if in_notes:
                    note_lines.append(text)

            elif child.tag == qn('w:tbl'):
                # A table ends the current notes block (notes are plain
                # paragraphs only, positioned between the inclusions table
                # and the next package heading)
                if in_notes:
                    _flush()
                    in_notes = False
        if in_notes:
            _flush()
    return "\n\n".join(blocks)


def extract_all_daywise_itinerary(docs_dir: Path) -> str:
    """
    Extract EVERY day-by-day itinerary block from every DOCX in docs_dir,
    tagged with its package heading, and return them concatenated as a
    single authoritative block — preserving route/distance segments in
    full and flagging optional/"upon request" activities SEPARATELY from
    guaranteed ones.

    WHY THIS EXISTS: same class of problem as price tables, hotel options,
    and important notes. A single day's itinerary content spans several
    consecutive paragraphs (a "Day N: <title>" heading, one or more
    route/distance parentheticals like "(Tbilisi – Gudauri: Distance: 120
    kms. / Driving Time: 2 hrs. 20 mins)", a long prose description, and an
    "Overnight in <city>" line) that must all be kept together AND
    correctly attributed to the right package. Chunk-based semantic
    retrieval frequently:
      - splits a day's route/distance segments away from its heading or
        drops them from the retrieved context entirely,
      - truncates or fragments the prose description, causing the LLM to
        compress a detailed multi-sentence day into a single generic
        bullet, and
      - fails to recognise an "upon request" / optional activity when the
        qualifying phrase ends up in a different chunk than the activity
        name, causing optional add-ons to be presented as guaranteed
        inclusions.

    FIX: bypass retrieval entirely for the day-wise itinerary. Walk each
    document in reading order; for every package heading, collect each
    "Day N: <title>" block verbatim — route/distance lines kept as their
    own list, the narrative split into sentences, and any sentence
    containing an optional/upon-request keyword pulled into a separate
    "Optional Activities Mentioned" list — up to the next Day heading,
    package heading, or a stop marker (pricing/accommodation/vehicle/
    inclusions/notes headings, or a table).
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    package_heading_re = re.compile(r'^\s*\d+\s*Nights?\b', re.I)
    day_re = re.compile(r'^\s*Day\s*(\d+)\s*[:\-]\s*(.*)$', re.I)
    overnight_re = re.compile(r'^\s*Overnight\s+in\s+(.+)$', re.I)
    route_re = re.compile(
        r'.*?(?:distance|driving\s*time|\bkms?\.?\b).*', re.I
    )
    stop_markers_re = re.compile(
        r'^\s*(Rates\s+in|INDEX|Accommodation\s+Details|Inclusions\s*&?\s*'
        r'Exclusions|Important\s*Notes|Vehicle\s+Details)', re.I
    )
    optional_kw_re = re.compile(
        r'upon\s+request|optional(?:\s+activity)?|available\s+on\s+request|'
        r'at\s+extra\s+cost|if\s+required|subject\s+to\s+availability|'
        r'can\s+be\s+added', re.I
    )

    def _split_optional_sentences(text: str) -> tuple[list, list]:
        """Split narrative into (regular_sentences, optional_sentences)
        based on presence of optional-activity keywords in each sentence."""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        regular, optional = [], []
        for s in sentences:
            s = s.strip()
            if not s:
                continue
            if optional_kw_re.search(s):
                optional.append(s)
            else:
                regular.append(s)
        return regular, optional

    blocks = []
    for f in sorted(Path(docs_dir).glob("*.docx")):
        try:
            doc = Document(f)
        except Exception:
            continue
        body = doc.element.body
        current_heading = ""
        current_day_num = None
        current_day_title = ""
        route_segments: list = []
        narrative_lines: list = []
        overnight_city = ""
        package_days: list = []

        def _flush_day():
            nonlocal current_day_num, current_day_title, route_segments
            nonlocal narrative_lines, overnight_city
            if current_day_num is None:
                return
            narrative = " ".join(narrative_lines).strip()
            regular, optional = _split_optional_sentences(narrative)
            day_block = [f"Day {current_day_num}: {current_day_title}"]
            if route_segments:
                day_block.append("  Route Segments:")
                for seg in route_segments:
                    day_block.append(f"    - {seg}")
            if regular:
                day_block.append("  Description:")
                for s in regular:
                    day_block.append(f"    {s}")
            if optional:
                day_block.append("  Optional Activities Mentioned (NOT guaranteed inclusions):")
                for s in optional:
                    day_block.append(f"    - {s}")
            if overnight_city:
                day_block.append(f"  Overnight: {overnight_city}")
            package_days.append("\n".join(day_block))
            current_day_num, current_day_title = None, ""
            route_segments, narrative_lines, overnight_city = [], [], ""

        def _flush_package():
            nonlocal package_days, current_heading
            _flush_day()
            if package_days and current_heading:
                blocks.append(
                    f"### DAY-WISE ITINERARY — Package: {current_heading} "
                    f"(source: {f.name})\n" + "\n\n".join(package_days)
                )
            package_days = []

        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                p = Paragraph(child, doc)
                text = p.text.strip()
                if not text:
                    continue

                if package_heading_re.match(text) and len(text) < 120:
                    _flush_package()
                    current_heading = text
                    continue

                day_m = day_re.match(text)
                if day_m:
                    _flush_day()
                    current_day_num = day_m.group(1)
                    current_day_title = day_m.group(2).strip()
                    continue

                if stop_markers_re.match(text):
                    _flush_day()
                    continue

                if current_day_num is not None:
                    ov_m = overnight_re.match(text)
                    if ov_m:
                        overnight_city = ov_m.group(1).strip()
                        continue
                    if route_re.match(text) and len(text) < 150:
                        route_segments.append(text.strip('() '))
                        continue
                    narrative_lines.append(text)

            elif child.tag == qn('w:tbl'):
                # A table always ends the current day — in this document
                # pricing/hotel/vehicle/inclusions tables always follow the
                # day-wise narrative, never interrupt it.
                _flush_day()

        _flush_package()  # flush the final package block in this file

    return "\n\n".join(blocks)


# ── Master Travel Plans (new, independent Excel-based knowledge source) ───────
# Isolated from the existing PDF/DOCX + ChromaDB knowledge base above. Reads
# docs/All_Plans_of_KUKUTRIP_Master.xlsx directly on every query (small file,
# loaded via st.cache_data so it is parsed once per session/file-mtime) and
# performs structured filtering instead of chunk-based vector similarity
# search, since the workbook's data is already normalized into tabular rows
# keyed by "Package ID" across sheets. This NEVER touches CHROMA_DIR, the
# existing `col` collection, or any of the docx/pdf extraction helpers above.
#
# Workbook structure (inspected directly — do not assume without re-checking
# if the file changes):
#   All 8 sheets share the same layout: 2 title rows + 1 blank row, then a
#   real header row at (0-based) index 3, then data rows. So every sheet
#   must be read with header=3.
#   - "Packages"   : Package ID, Package Name, Country, Nights, Days, Route,
#                    Validity, Hotel Options, From (USD pp)*  — one row per
#                    package (204 rows spanning many countries/combinations)
#   - "Pricing"    : Package ID, Package Name, Nights, Hotel Option, Category
#                    (star rating e.g. '5*', '4*/3*'), Pax Tier (e.g. '2 Pax',
#                    '4 Pax (USD)'), Price / Person (USD), Single Supp (USD),
#                    Child No-Bed 0-11 (USD) — many rows per package (one per
#                    hotel-category × pax-tier combination)
#   - "Itinerary"  : Package ID, Package Name, Day, Title, Highlights — one
#                    row per day per package
#   - "Inclusions" : Package ID, Package Name, Common Inclusions,
#                    Package-Specific Entries, Exclusions, Important Notes —
#                    one row per package
#   - "Travel Guides & Country Info": Destination / Country, Topic / Category,
#                    Key Guidelines & Critical Advisories, Permit/Visa/ID
#                    Requirements, Currency & Local Regulations
#   - "Experiences & Activities": Destination / City, Experience / Activity
#                    Name, Category / Type, Indicative Cost / Pricing,
#                    Description & PDF Highlight Bullet, Operational Notes
#   - "Terms & Conditions": Destination / Country, Policy Domain / Subject,
#                    Standard Clause / Policy Details, Deadlines / Financial
#                    Penalities, Operational Notes for PDF Generation
#   - "Company Info": Attribute / Field, Official Details, Usage in Proposal

MASTER_PLANS_SHEET_HEADER_ROW = 3  # 0-based header row index, same for all sheets


@st.cache_data(show_spinner="Loading KukuTrip Master Travel Plans…")
def load_master_plans_workbook(_mtime: float):
    """
    Load every sheet of the Master Travel Plans workbook into DataFrames.
    `_mtime` (file modification time) is passed purely to bust the
    st.cache_data cache when the source .xlsx is replaced/updated. Returns
    a dict of {sheet_name: DataFrame}, or None if the file is missing/
    unreadable — callers must handle that explicitly.
    """
    if not MASTER_PLANS_XLSX.exists():
        return None
    try:
        import pandas as pd
        xls = pd.ExcelFile(MASTER_PLANS_XLSX)
        sheets = {}
        for name in xls.sheet_names:
            df = xls.parse(name, header=MASTER_PLANS_SHEET_HEADER_ROW)
            df = df.dropna(how="all")  # drop spacer rows
            sheets[name] = df
        return sheets
    except Exception:
        return None


def get_master_plans_workbook():
    """Cache-aware accessor — re-reads the workbook only if its mtime changes."""
    if not MASTER_PLANS_XLSX.exists():
        return None
    mtime = MASTER_PLANS_XLSX.stat().st_mtime
    return load_master_plans_workbook(mtime)


def _norm_text(s) -> str:
    return str(s or "").strip().lower()

def master_plans_all_countries(sheets: dict) -> list:
    """Distinct country/route names from the Packages sheet, for error messages."""
    df = sheets.get("Packages")
    if df is None or "Country" not in df.columns:
        return []
    return sorted(df["Country"].dropna().astype(str).unique().tolist())


def search_master_plans(sheets: dict, destination: str = "", nights=None, days=None):
    """
    Filter the Packages sheet for rows matching the given destination
    (substring match against Country, case-insensitive — matches combined
    routes like "Kazakhstan & Georgia" when searching "Georgia") and/or an
    exact Nights or Days count. Returns the matching rows (may be empty).
    Does not invent or fall back to unrelated packages — only real rows
    from the sheet are ever returned.
    """
    import pandas as pd
    df = sheets.get("Packages")
    if df is None:
        return pd.DataFrame()
    out = df.copy()
    if destination:
        needle = _norm_text(destination)
        out = out[out["Country"].astype(str).str.lower().str.contains(needle, na=False)]
    if nights is not None:
        out = out[pd.to_numeric(out["Nights"], errors="coerce") == nights]
    if days is not None:
        out = out[pd.to_numeric(out["Days"], errors="coerce") == days]
    return out


def get_master_plan_itinerary(sheets: dict, package_id: str):
    import pandas as pd
    df = sheets.get("Itinerary")
    if df is None:
        return pd.DataFrame()
    return df[df["Package ID"].astype(str) == str(package_id)].sort_values(
        by="Day", key=lambda s: pd.to_numeric(s, errors="coerce")
    )


def get_master_plan_inclusions(sheets: dict, package_id: str) -> dict:
    df = sheets.get("Inclusions")
    if df is None:
        return {}
    rows = df[df["Package ID"].astype(str) == str(package_id)]
    if rows.empty:
        return {}
    r = rows.iloc[0]
    return {
        "common_inclusions": str(r.get("Common Inclusions", "") or ""),
        "package_specific": str(r.get("Package-Specific Entries", "") or ""),
        "exclusions": str(r.get("Exclusions", "") or ""),
        "important_notes": str(r.get("Important Notes", "") or ""),
    }


def get_master_plan_pricing(sheets: dict, package_id: str):
    df = sheets.get("Pricing")
    if df is None:
        return df
    return df[df["Package ID"].astype(str) == str(package_id)]


def get_master_plan_country_info(sheets: dict, country: str):
    import pandas as pd
    df = sheets.get("Travel Guides & Country Info")
    if df is None:
        return pd.DataFrame()
    needle = _norm_text(country)
    return df[df["Destination / Country"].astype(str).str.lower().str.contains(needle, na=False)]


def get_master_plan_experiences(sheets: dict, destination: str):
    import pandas as pd
    df = sheets.get("Experiences & Activities")
    if df is None:
        return pd.DataFrame()
    needle = _norm_text(destination)
    return df[df["Destination / City"].astype(str).str.lower().str.contains(needle, na=False)]


def get_master_plan_terms(sheets: dict, country: str):
    import pandas as pd
    df = sheets.get("Terms & Conditions")
    if df is None:
        return pd.DataFrame()
    needle = _norm_text(country)
    return df[df["Destination / Country"].astype(str).str.lower().str.contains(needle, na=False)]


def build_master_plans_context(sheets: dict, matched_packages) -> str:
    """
    Build a compact, LLM-ready context block containing ONLY the rows
    relevant to the matched package(s) — package summary, full day-wise
    itinerary, inclusions/exclusions, pricing grid, and (if a small number
    of distinct countries are involved) country guide/terms/experiences
    info. Mirrors the existing docx pipeline's "authoritative block"
    pattern (bypass vector retrieval, inject exact structured rows) but
    reads from the Excel sheets instead of docx paragraphs/tables.
    """
    if matched_packages is None or matched_packages.empty:
        return ""

    blocks = []
    countries_seen = set()
    for _, pkg in matched_packages.iterrows():
        pid = str(pkg.get("Package ID", ""))
        if not pid:
            continue
        countries_seen.add(str(pkg.get("Country", "")))

        blocks.append(
            f"### PACKAGE SUMMARY — {pid}\n"
            f"Package Name: {pkg.get('Package Name','')}\n"
            f"Country: {pkg.get('Country','')}\n"
            f"Nights: {pkg.get('Nights','')} | Days: {pkg.get('Days','')}\n"
            f"Route: {pkg.get('Route','')}\n"
            f"Validity: {pkg.get('Validity','')}\n"
            f"Hotel Options available: {pkg.get('Hotel Options','')}\n"
            f"Indicative starting price: USD {pkg.get('From (USD pp)*','')} per person"
        )

        itin = get_master_plan_itinerary(sheets, pid)
        if not itin.empty:
            lines = [f"### DAY-WISE ITINERARY — {pid} ({pkg.get('Package Name','')})"]
            for _, day in itin.iterrows():
                lines.append(
                    f"Day {day.get('Day','')}: {day.get('Title','')} — {day.get('Highlights','')}"
                )
            blocks.append("\n".join(lines))

        incl = get_master_plan_inclusions(sheets, pid)
        if incl:
            blocks.append(
                f"### INCLUSIONS & EXCLUSIONS — {pid}\n"
                f"Common Inclusions: {incl['common_inclusions']}\n"
                f"Package-Specific Entries: {incl['package_specific']}\n"
                f"Exclusions: {incl['exclusions']}\n"
                f"Important Notes: {incl['important_notes']}"
            )

        pricing = get_master_plan_pricing(sheets, pid)
        if pricing is not None and not pricing.empty:
            lines = [f"### PRICE GRID (per person, twin-sharing, USD) — {pid}"]
            for _, row in pricing.iterrows():
                lines.append(
                    f"Hotel Option: {row.get('Hotel Option','')} | Category: {row.get('Category','')} "
                    f"| Pax Tier: {row.get('Pax Tier','')} | Price/Person (USD): {row.get('Price / Person (USD)','')} "
                    f"| Single Supp (USD): {row.get('Single Supp (USD)','')} "
                    f"| Child No-Bed 0-11 (USD): {row.get('Child No-Bed 0-11 (USD)','')}"
                )
            blocks.append("\n".join(lines))

    if 0 < len(countries_seen) <= 3:
        for country in countries_seen:
            if not country:
                continue
            info = get_master_plan_country_info(sheets, country)
            if not info.empty:
                lines = [f"### TRAVEL GUIDE & COUNTRY INFO — {country}"]
                for _, row in info.iterrows():
                    lines.append(
                        f"{row.get('Topic / Category','')}: "
                        f"{row.get('Key Guidelines & Critical Advisories','')} "
                        f"| Permit/Visa: {row.get('Permit / Visa / ID Requirements','')} "
                        f"| Currency: {row.get('Currency & Local Regulations','')}"
                    )
                blocks.append("\n".join(lines))

            terms = get_master_plan_terms(sheets, country)
            if not terms.empty:
                lines = [f"### TERMS & CONDITIONS — {country}"]
                for _, row in terms.iterrows():
                    lines.append(
                        f"{row.get('Policy Domain / Subject','')}: "
                        f"{row.get('Standard Clause / Policy Details','')} "
                        f"({row.get('Deadlines / Financial Penalities','')})"
                    )
                blocks.append("\n".join(lines))

            exp = get_master_plan_experiences(sheets, country)
            if not exp.empty:
                lines = [f"### OPTIONAL EXPERIENCES & ACTIVITIES — {country}"]
                for _, row in exp.iterrows():
                    lines.append(
                        f"{row.get('Experience / Activity Name','')} "
                        f"({row.get('Destination / City','')}, {row.get('Category / Type','')}): "
                        f"{row.get('Description & PDF Highlight Bullet','')} "
                        f"[{row.get('Indicative Cost / Pricing','')}]"
                    )
                blocks.append("\n".join(lines))

    return "\n\n".join(blocks)

def get_master_plans_answer(question: str, history: list, api_key: str, model: str,
                            destination_hint: str = "", nights_hint=None,
                            days_hint=None):
    """
    Master Travel Plans equivalent of get_answer() — same clarification /
    JSON-itinerary-output contract (so the rest of the app — pending_meta,
    pricing workflow, PDF generation — works unmodified), but retrieves
    from the Excel workbook via structured filtering instead of ChromaDB
    vector search. Completely isolated from `col`/build_vectorstore/
    get_answer above; the existing knowledge base is never read here.
    Returns (answer_markdown, sources_list, meta_dict).
    """
    sheets = get_master_plans_workbook()
    if sheets is None:
        return (
            "⚠️ The KukuTrip Master Travel Plans file "
            f"(`{MASTER_PLANS_XLSX.name}`) could not be found or read under `docs/`. "
            "Please make sure it has been uploaded, then try again.",
            [], {}
        )

    matched = search_master_plans(sheets, destination=destination_hint,
                                  nights=nights_hint, days=days_hint)
    if matched.empty and destination_hint:
        matched_dest_only = search_master_plans(sheets, destination=destination_hint)
        if matched_dest_only.empty:
            available = master_plans_all_countries(sheets)
            return (
                f"I couldn't find a matching plan in the selected **KukuTrip Master "
                f"Travel Plans** source for **{destination_hint}**. Available "
                f"destinations include: {', '.join(available[:25])}"
                f"{'…' if len(available) > 25 else ''}. Please choose a different "
                "destination, or check the spelling.",
                [], {}
            )
        matched = matched_dest_only

    # ── Avoid dumping the ENTIRE workbook into the LLM context ──────────────
    # If no destination was recognised at all, `search_master_plans` returns
    # every package (200+ rows across 46 countries). Building the full
    # itinerary/pricing/inclusions context for all of them makes the prompt
    # enormous, which is what was causing the Gemini API call to take a long
    # time and eventually hit the 90s read timeout. Instead, short-circuit
    # with a cheap, instant clarification question (no LLM call at all) so
    # the user picks a destination first — mirroring the "act like a travel
    # consultant, ask before generating" requirement.
    MAX_PACKAGES_FOR_LLM_CONTEXT = 12
    if not destination_hint and len(matched) > MAX_PACKAGES_FOR_LLM_CONTEXT:
        available = master_plans_all_countries(sheets)
        return (
            "I'd love to help plan your trip! Could you tell me which "
            "**destination/country** you're interested in? Available options "
            f"in our KukuTrip Master Travel Plans include: {', '.join(available[:25])}"
            f"{'…' if len(available) > 25 else ''}.",
            [], {}
        )
    # Even with a destination match, cap how many package rows are expanded
    # into full context (summary + itinerary + pricing + inclusions each) to
    # keep the prompt size — and therefore response time — reasonable. The
    # cap is generous enough to cover every real destination/nights
    # combination in the workbook, which rarely exceeds a handful of exact
    # matches.
    if len(matched) > MAX_PACKAGES_FOR_LLM_CONTEXT:
        matched = matched.head(MAX_PACKAGES_FOR_LLM_CONTEXT)

    sources = sorted(set(matched["Package ID"].astype(str).tolist())) if not matched.empty else []
    context = build_master_plans_context(sheets, matched)

    history_text = "".join(
        f"{'User' if m['role']=='user' else 'Assistant'}: {m['content']}\n"
        for m in history[-6:]
    )

    system = (
        "You are a professional Travel Itinerary Generation Agent acting as a "
        "knowledgeable travel consultant.\n\n"

        "## KNOWLEDGE SOURCE — SOURCE OF TRUTH\n"
        "The 'KukuTrip Master Travel Plans' data provided below (extracted from the "
        "master Excel database of pre-built packages) is your ONLY source for: "
        "destinations, package names, nights/days, routes, hotels, day-wise "
        "itineraries, inclusions, exclusions, important notes, and pricing. "
        "NEVER invent travel information not present in the data below.\n\n"

        "## CLARIFICATION WORKFLOW (FOLLOW EXACTLY — ACT LIKE A TRAVEL CONSULTANT)\n"
        "STEP 1 — Understand the request.\n"
        "STEP 2 — Check if critical information is missing:\n"
        "  • Destination/country — if missing, ASK\n"
        "  • Travel dates — ask if needed for the itinerary/pricing\n"
        "  • Number of nights/days — if missing and multiple packages of "
        "different durations exist for the destination, ASK\n"
        "  • Number of adults, children (with ages if relevant), and rooms — if "
        "missing and pricing/hotel selection depends on it, ASK\n"
        "  • Preferred transportation — if relevant options exist, ASK\n"
        "  • Hotel/category preference — if the matched package(s) below list "
        "multiple hotel/star options, ASK which one\n"
        "STEP 3 — If multiple materially different packages match (different "
        "routes, durations, or hotel categories) — SHOW OPTIONS, wait for the "
        "user's choice.\n"
        "STEP 4 — Combine all missing questions into ONE concise message (max 4 "
        "questions at once). Do not ask unnecessary questions — if only one "
        "package/option exists below, use it automatically.\n"
        "STEP 5 — Only once enough information is available AND (if applicable) "
        "the user has picked a specific package/hotel category, respond in TWO "
        "parts separated by ---READY---\n\n"
        "PART 1: Friendly Markdown itinerary summary ending with:\n"
        "'✅ **Your itinerary is ready. Click Generate PDF to create your document.**'\n\n"
        "---READY---\n\n"
        "PART 2: Structured JSON (schema below). Do NOT output ---READY--- during "
        "## JSON SCHEMA\n"
        "{\n"
        '  "destination": "Country name",\n'
        '  "package_name": "Exact or adapted package name from the data",\n'
        '  "image_keyword": "Most iconic landmark/city for this destination",\n'
        '  "route": "City1 · City2 · City3",\n'
        '  "dates": "DD MMM – DD MMM YYYY (or empty if not provided)",\n'
        '  "nights": "N Nights / N Days",\n'
        '  "persons": "N Adults · N Children · N Rooms",\n'
        '  "transport": "e.g. Private Cab, Self Drive",\n'
        '  "days": [\n'
        '    {"day": 1, "date": "DD MMM or empty", "title": "Day title",\n'
        '     "activities": ["Activity 1", "Activity 2"],\n'
        '     "overnight": "City (omit on final/departure day)"}\n'
        "  ],\n"
        '  "hotels": [{"city": "City (N Nights)", "hotel": "Hotel — Meal Plan", "dates": "DD–DD MMM"}],\n'
        '  "highlights": ["Highlight 1"],\n'
        '  "inclusions": ["Only applicable items"],\n'
        '  "exclusions": ["Only applicable items"],\n'
        '  "notes": ["Only applicable notes"],\n'
        '  "amount": "Price from data e.g. USD 439 per person — omit if not available"\n'
        "}\n\n"

        "## DAY-WISE ITINERARY\n"
        "Use the 'Title' and 'Highlights' fields from the DAY-WISE ITINERARY block "
        "below verbatim for each day — split the Highlights text on ';' or ',' into "
        "separate 'activities' array items rather than one long string. Number days "
        "sequentially exactly as given. Infer the 'overnight' city from the route/"
        "day title context (e.g. a day titled 'Tbilisi to Batumi' ends overnight in "
        "Batumi) — never fabricate a city not mentioned in the route.\n\n"

        "## PRICING\n"
        "The PRICE GRID block lists price-per-person in USD by hotel category and "
        "pax tier. Match the user's confirmed hotel category and traveller count to "
        "select the correct row; if the user hasn't specified either yet, ASK before "
        "quoting a specific price. Never invent a price not present in the grid.\n\n"

        "## INCLUSIONS/EXCLUSIONS/NOTES\n"
        "Copy the Common Inclusions and Package-Specific Entries (combined) "
        "verbatim into the JSON 'inclusions' array (split on ';'), Exclusions into "
        "'exclusions' (split on ';'), and Important Notes into 'notes' (as one or "
        "more items) — do not omit or paraphrase them.\n\n"

        "## PAX / PERSON / ADULT NORMALIZATION\n"
        "Pax, Persons, Adults all refer to the same field (number of travellers). "
        "Always output the JSON 'persons' field as 'N Adults · N Children · N Rooms'.\n\n"

        "## DO NOT INVENT MISSING DATA\n"
        "If the matched package data below does not include something (e.g. no "
        "exact hotel name for a requested category), say so and ask the user, "
        "rather than fabricating it. Distinguish clearly between user-provided "
        "information, data from the Master Travel Plans source, and any "
        "estimate/assumption (which must be explicitly marked as such if used).\n\n"

        f"## KUKUTRIP MASTER TRAVEL PLANS DATA (matched to this request)\n{context}\n\n"
        f"## CONVERSATION HISTORY\n{history_text}"
    )

    # Larger timeout/attempts than the default: the injected context here
    # (package summary + full day-wise itinerary + price grid + inclusions
    # for up to MAX_PACKAGES_FOR_LLM_CONTEXT packages) is typically bigger
    # than the docx-based Knowledge Base context, so generation can
    # legitimately take longer.
    raw = call_gemini(api_key, model, system, question, timeout=150, max_attempts=2)

    meta = {}
    answer = raw
    for sep in ("---READY---", "---JSON---"):
        if sep in raw:
            parts = raw.split(sep, 1)
            answer = parts[0].strip()
            json_str = parts[1].strip()
            json_str = re.sub(r'^```[a-z]*\n?', '', json_str).strip()
            json_str = re.sub(r'\n?```$', '', json_str).strip()
            try:
                meta = json.loads(json_str)
            except Exception:
                m = re.search(r'\{.*\}', json_str, re.DOTALL)
                if m:
                    try:
                        meta = json.loads(m.group())
                    except Exception:
                        pass
            break

    if not meta:
        m = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', raw, re.DOTALL)
        if not m:
            m = re.search(r'(\{\s*"destination".*?"notes".*?\})', raw, re.DOTALL)
        if m:
            try:
                meta = json.loads(m.group(1))
                answer = raw[:m.start()].strip()
                if not answer:
                    answer = f"Here is your itinerary for **{meta.get('destination','')}**."
            except Exception:
                pass

    return answer, sources, meta


def extract_master_plans_hints(text: str) -> dict:
    """
    Lightweight heuristic extraction of destination/nights/days hints from
    the user's latest free-text message, used to pre-filter the Packages
    sheet before invoking the LLM. This is best-effort only — the LLM's
    own clarification workflow (in get_master_plans_answer's system
    prompt) is the authoritative fallback for anything this misses; it
    never blocks the conversation, it only narrows the initial data slice
    handed to the model.
    """
    hints = {"destination": "", "nights": None, "days": None}
    m = re.search(r'(\d+)\s*[Nn]ight', text)
    if m:
        hints["nights"] = int(m.group(1))
    m = re.search(r'(\d+)\s*[Dd]ay', text)
    if m:
        hints["days"] = int(m.group(1))
    return hints


def guess_master_plans_destination(sheets: dict, text: str) -> str:
    """
    Best-effort match of a country/destination name mentioned in free text
    against the actual set of countries present in the Packages sheet
    (including combo routes like "Kazakhstan & Georgia"). Only ever
    returns a name that literally exists in the sheet — never invents or
    guesses beyond what's really there. Returns "" if nothing matches.
    """
    countries = master_plans_all_countries(sheets)
    needle = _norm_text(text)
    # Prefer the longest matching country/combo name to avoid a generic
    # single-country match masking a more specific combo match.
    matches = [c for c in countries if c and _norm_text(c) in needle]
    if matches:
        return max(matches, key=len)
    # Fall back to checking individual words of each combo against the text
    for c in sorted(countries, key=len, reverse=True):
        parts = re.split(r'\s*&\s*|,\s*', c)
        for p in parts:
            if p and _norm_text(p) in needle:
                return c
    return ""



# ── Vector store ──────────────────────────────────────────────────────────────





@st.cache_resource(show_spinner="Building document index…")
def build_vectorstore():
    import chromadb
    from chromadb.utils.embedding_functions import DefaultEmbeddingFunction
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    ef     = DefaultEmbeddingFunction()
    client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    name   = "travel_docs"

    try:
        col = client.get_collection(name, embedding_function=ef)
        if col.count() > 0:
            return col
        client.delete_collection(name)
    except Exception:
        pass

    DOCS_DIR.mkdir(exist_ok=True)
    raw = []
    for f in sorted(DOCS_DIR.iterdir()):
        try:
            if f.suffix.lower() == ".pdf":
                raw.append({"source": f.name, "content": load_pdf_text(f)})
            elif f.suffix.lower() == ".docx":
                raw.append({"source": f.name, "content": load_docx_text(f)})
        except Exception as e:
            st.warning(f"Could not load {f.name}: {e}")

    if not raw:
        return None

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    ids, docs, metas = [], [], []
    for d in raw:
        for i, chunk in enumerate(splitter.split_text(d["content"])):
            ids.append(f"{d['source']}_{i}")
            docs.append(chunk)
            metas.append({"source": d["source"]})

    col = client.create_collection(name, embedding_function=ef)
    for start in range(0, len(docs), 100):
        col.add(ids=ids[start:start+100], documents=docs[start:start+100], metadatas=metas[start:start+100])
    return col


# ── Gemini call ───────────────────────────────────────────────────────────────

def call_gemini(api_key: str, model: str, system: str, user: str,
                timeout: int = 120, max_attempts: int = 3) -> str:
    """
    Call the Gemini generateContent API, with automatic retry/backoff on
    HTTP 429 (rate limit / quota exceeded), 503 (transiently overloaded),
    and network-level read timeouts. Gemini's free tier enforces both a
    requests-per-minute and a requests-per-day quota; a 429 almost always
    means one of those limits was hit. Large prompts (e.g. many matched
    packages injected as context) can also legitimately take longer than
    the default timeout to generate a full JSON itinerary, so timeouts are
    retried too rather than failing immediately. Gives up with a clear,
    actionable error message instead of a raw HTTPError/ReadTimeout
    traceback.

    `timeout` (seconds) and `max_attempts` are tunable per-call so callers
    with especially large contexts (e.g. Master Travel Plans) can allow
    more time/attempts without changing every other caller's behavior.
    """
    import requests, time
    from requests.exceptions import ReadTimeout, ConnectionError as ReqConnectionError

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    _p = os.environ.get("HTTPS_PROXY") or os.environ.get("HTTP_PROXY") or ""
    proxies = {"https": _p, "http": _p} if _p else {}
    payload = {
        "system_instruction": {"parts": [{"text": system}]},
        "contents": [{"role": "user", "parts": [{"text": user}]}],
        # NOTE: maxOutputTokens must be generous enough to fit the FULL JSON
        # itinerary — including the verbatim, per-package inclusions/exclusions
        # (now up to ~20 items each, per extract_all_inclusion_exclusion_tables)
        # PLUS the day-wise itinerary, hotels, highlights, notes, and amount
        # fields that come after inclusions/exclusions in the JSON schema.
        # A too-small limit here silently truncates the JSON near the end
        # (notes/amount), which then disappears from the generated PDF
        # even though earlier sections still render correctly.
        "generationConfig": {"temperature": 0.7, "maxOutputTokens": 8192}
    }

    last_err = None
    last_timeout = False
    for attempt in range(1, max_attempts + 1):
        try:
            resp = requests.post(url, json=payload, headers={"Content-Type": "application/json"},
                                 params={"key": api_key}, proxies=proxies, timeout=timeout)
        except (ReadTimeout, ReqConnectionError) as e:
            last_err = e
            last_timeout = True
            if attempt == max_attempts:
                break
            time.sleep(min(10, 2 ** attempt))
            continue

        if resp.status_code in (429, 503):
            last_err = resp
            last_timeout = False
            if attempt == max_attempts:
                break
            # Prefer the server's suggested retry delay if present in the
            # error body (Gemini returns RetryInfo.retryDelay e.g. "23s").
            wait_s = min(30, 2 ** attempt)
            try:
                body = resp.json()
                for detail in body.get("error", {}).get("details", []):
                    rd = detail.get("retryDelay")
                    if rd:
                        wait_s = min(30, float(str(rd).rstrip("s")) + 1)
                        break
            except Exception:
                pass
            time.sleep(wait_s)
            continue
        resp.raise_for_status()
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"]

    # Exhausted all retries.
    if last_timeout:
        raise RuntimeError(
            f"Gemini API did not respond within {timeout}s after {max_attempts} "
            "attempts. This can happen when the request/context is very large "
            "(e.g. too many matched travel packages) or the network/proxy is "
            "slow. Please try narrowing your request (e.g. a specific "
            "destination and duration) and try again."
        )
    if last_err is not None and getattr(last_err, "status_code", None) == 429:
        raise RuntimeError(
            "Gemini API rate limit / quota exceeded (HTTP 429) after several "
            "retries. This usually means the free-tier requests-per-minute "
            "or requests-per-day limit for this API key/model has been hit. "
            "Please wait a minute and try again, switch to a different model "
            "in the sidebar, or use an API key with a higher quota."
        )
    if last_err is not None:
        status = getattr(last_err, "status_code", "unknown")
        raise RuntimeError(
            f"Gemini API returned HTTP {status} after several "
            "retries — the service may be temporarily overloaded. Please try again shortly."
        )
    raise RuntimeError("Gemini API call failed for an unknown reason.")


def get_answer(col, question: str, history: list, api_key: str, model: str) -> tuple[str, list[str], dict]:
    # Primary query
    results = col.query(query_texts=[question], n_results=12)
    docs_seen = set(results["documents"][0])
    all_docs = list(results["documents"][0])
    sources = list({m["source"] for m in results["metadatas"][0]})

    # Always include a secondary hotel-specific retrieval so hotel tables are always in context
    hotel_query = "hotel accommodation destination star category 5 star 4 star 3 star"
    h_results = col.query(query_texts=[hotel_query], n_results=8)
    for d in h_results["documents"][0]:
        if d not in docs_seen:
            all_docs.append(d)
            docs_seen.add(d)
    for m in h_results["metadatas"][0]:
        sources.append(m["source"])
    sources = list(set(sources))

    context = "\n\n".join(all_docs)

    # ── Always inject the COMPLETE, authoritative price tables ──────────────
    # Semantic retrieval alone is unreliable for pricing data: when a
    # document contains multiple near-identical price tables (same headers,
    # same "2 Pax / 4 Pax / 6 Pax" row labels, differing only in the numeric
    # values), vector similarity search can retrieve a fragment or the WRONG
    # package's table into context, causing the LLM to quote incorrect
    # rates. To guarantee correctness, bypass retrieval entirely for price
    # tables and inject the full, correctly-labeled set directly.
    price_tables_block = extract_all_price_tables(DOCS_DIR)
    if price_tables_block:
        context += (
            "\n\n## AUTHORITATIVE PRICE TABLES (use these EXACTLY — "
            "do not use price figures found elsewhere in the context above; "
            "each table is labeled with its exact package name — match the "
            "package name precisely before using its rates)\n"
            + price_tables_block
        )

    # ── Always inject the COMPLETE, authoritative inclusions/exclusions ─────
    # Same rationale as price tables: each package's nights/city-specific
    # inclusion & exclusion list (e.g. "3 Nights' accommodation in Tbilisi",
    # "1 Night accommodation in Gudauri") is easily confused with a
    # different package's list by semantic retrieval, since the wording
    # overlaps heavily ("Daily Breakfast", "Round Airport Transfers", etc.
    # appear in nearly every package). Bypass retrieval and inject the full,
    # correctly-labeled set directly so the model always matches the exact
    # package name before using its inclusions/exclusions.
    incl_excl_block = extract_all_inclusion_exclusion_tables(DOCS_DIR)
    if incl_excl_block:
        context += (
            "\n\n## AUTHORITATIVE INCLUSIONS & EXCLUSIONS (use these EXACTLY — "
            "do not use inclusion/exclusion items found elsewhere in the context "
            "above; each block is labeled with its exact package name — match "
            "the package name/nights-breakdown precisely (e.g. '5 Nights - 4 "
            "Nights Tbilisi + 1 Night Gudauri' is a DIFFERENT package from '6 "
            "Nights - 4 Nights Tbilisi + 2 Nights Batumi') before using its "
            "inclusions/exclusions — do NOT mix items from a different "
            "package's list even if the wording looks similar)\n"
            + incl_excl_block
        )

    # ── Always inject the COMPLETE, authoritative hotel options ─────────────
    # Same rationale as price tables and inclusions/exclusions: each
    # package's accommodation table lists city-specific hotel options per
    # star category, but the table structure/wording is nearly identical
    # across packages (differing mainly in which cities appear). Bypass
    # retrieval and inject the full, correctly-labeled set directly.
    hotel_tables_block = extract_all_hotel_tables(DOCS_DIR)
    if hotel_tables_block:
        context += (
            "\n\n## AUTHORITATIVE HOTEL OPTIONS (use these EXACTLY — "
            "do not use hotel names found elsewhere in the context above; "
            "each table is labeled with its exact package name — match the "
            "package name precisely, then pick the hotel option(s) for the "
            "correct destination/city and the user's requested star "
            "category from that table)\n"
            + hotel_tables_block
        )

    # ── Always inject the COMPLETE, authoritative Important Notes ──────────
    # "Important Notes" is free-text (not a table) and was previously left
    # to unreliable semantic retrieval, causing it to be dropped entirely
    # from many generated itineraries. Bypass retrieval and inject the full,
    # correctly-labeled set directly, per package.
    notes_block = extract_all_important_notes(DOCS_DIR)
    if notes_block:
        context += (
            "\n\n## AUTHORITATIVE IMPORTANT NOTES (use these EXACTLY — "
            "do not use notes found elsewhere in the context above; each "
            "block is labeled with its exact package name — match the "
            "package name precisely before copying its notes verbatim into "
            "the JSON 'notes' array)\n"
            + notes_block
        )

    # ── Always inject the COMPLETE, authoritative day-wise itinerary ────────
    # Semantic retrieval is especially unreliable for the day-wise narrative:
    # each day's route/distance segments, long prose description, and any
    # "upon request"/optional activity phrase must stay together AND be
    # correctly attributed to the right package — but with 6 near-identical
    # multi-day itineraries in one document, chunking frequently splits a
    # day's route segments away from its description, truncates the prose
    # into an over-short bullet, or loses the distinction between a
    # guaranteed and an "upon request" activity. Bypass retrieval and inject
    # the full, correctly-labeled, sentence-preserving day-wise breakdown
    # directly, per package.
    daywise_block = extract_all_daywise_itinerary(DOCS_DIR)
    if daywise_block:
        context += (
            "\n\n## AUTHORITATIVE DAY-WISE ITINERARY (use these EXACTLY — "
            "do not shorten or paraphrase the Description sentences into a "
            "single generic bullet; preserve every Route Segment (distance/ "
            "driving time) as its own bullet; each block is labeled with its "
            "exact package name — match the package name precisely, then "
            "use ALL of that package's Route Segments and Description "
            "sentences as separate 'activities' array items for that day; "
            "items listed under 'Optional Activities Mentioned' are NOT "
            "guaranteed inclusions — mark them clearly as optional/upon "
            "request in the day's activities, never present them as if "
            "already included)\n"
            + daywise_block
        )



    history_text = "".join(


        f"{'User' if m['role']=='user' else 'Assistant'}: {m['content']}\n"
        for m in history[-6:]
    )

    system = (
        "You are a professional Travel Itinerary Generation Agent.\n\n"

        "## KNOWLEDGE BASE — SOURCE OF TRUTH\n"
        "The documents provided are your ONLY source for: hotels, activities, prices, inclusions, "
        "exclusions, transfers, tours, and company policies. NEVER invent travel information.\n\n"

        "## CLARIFICATION WORKFLOW (FOLLOW EXACTLY)\n\n"
        "STEP 1 — Understand the request.\n"
        "STEP 2 — Search the knowledge base.\n"
        "STEP 3 — Check if critical information is missing:\n"
        "  • Country / destination — if missing, ASK\n"
        "  • Duration (days/nights) — if missing and cannot be inferred, ASK\n"
        "  • Traveller count — if missing and affects hotels/pricing, ASK\n"
        "  • Dates — only ask if they are needed for hotel assignment or pricing\n"
        "STEP 4 — Check if multiple materially different options exist in the KB:\n"
        "  • Multiple itinerary routes → SHOW OPTIONS, wait for user choice\n"
        "  • Multiple hotels for same city → SHOW OPTIONS, wait for user choice\n"
        "  • Multiple transportation modes → SHOW OPTIONS, wait for user choice\n"
        "  • Multiple packages with different durations/prices → SHOW OPTIONS\n"
        "STEP 5 — If info missing OR options exist: respond with Markdown ONLY (no ---JSON---).\n"
        "STEP 6 — When ALL information is confirmed: generate full itinerary + ---JSON--- block.\n\n"

        "## SMART DEFAULTS (DO NOT ASK UNNECESSARILY)\n"
        "- If only ONE option exists in KB → use it automatically\n"
        "- If user's request exactly matches one option → use it automatically\n"
        "- If user says 'you choose' or 'best option' → pick and explain, then generate\n"
        "- Do NOT ask for budget unless needed to filter between options\n"
        "- Do NOT ask for rooms if not relevant to the itinerary\n"
        "- Combine all missing questions into ONE message (max 3 questions at once)\n\n"

        "## OPTION PRESENTATION FORMAT\n"
        "When presenting options, use numbered list:\n"
        "**1.** Rome + Florence + Venice — 7 Days\n"
        "**2.** Rome + Florence + Dolomites — 7 Days\n"
        "Then ask: 'Which option would you prefer? Also let me know your travel dates and traveller count.'\n\n"

        "## WHEN READY — RESPONSE FORMAT\n"
        "Only when all choices are confirmed, respond in TWO parts separated by ---READY---\n\n"
        "PART 1: Friendly Markdown itinerary summary ending with:\n"
        "'✅ **Your itinerary is ready. Click Generate PDF to create your document.**'\n\n"
        "---READY---\n\n"
        "PART 2: Structured JSON (schema below). Do NOT output ---READY--- during clarification.\n\n"

        "## JSON SCHEMA\n"
        "{\n"
        '  "destination": "Country name",\n'
        '  "package_name": "Creative title e.g. GEORGIA WITH CAUCASUS MOUNTAINS",\n'
        '  "image_keyword": "Most iconic landmark e.g. Dolomites, Gergeti Church Kazbegi, Colosseum Rome",\n'
        '  "route": "City1 · City2 · City3",\n'
        '  "dates": "DD MMM – DD MMM YYYY (or empty if not provided)",\n'
        '  "nights": "N Nights / N Days",\n'
        '  "persons": "N Adults · N Children · N Rooms",\n'
        '  "transport": "e.g. Private Cab, Self Drive",\n'
        '  "days": [\n'
        '    {"day": 1, "date": "DD MMM or empty", "title": "Day title",\n'
        '     "activities": ["Activity 1", "Activity 2"],\n'
        '     "overnight": "City (omit on final/departure day)"}\n'
        "  ],\n"
        '  "hotels": [{"city": "City (N Nights)", "hotel": "Hotel — Meal Plan", "dates": "DD–DD MMM"}],\n'
        '  "highlights": ["Highlight 1"],\n'
        '  "inclusions": ["Only applicable items"],\n'
        '  "exclusions": ["Only applicable items"],\n'
        '  "notes": ["Only applicable notes"],\n'
        '  "amount": "Price from KB e.g. INR 1,45,000 per person — omit if not in KB"\n'
        "}\n\n"

        "## DAY-WISE ITINERARY DETAIL REQUIREMENT (MANDATORY)\n"
        "Each day's 'activities' array MUST preserve the full descriptive detail from the "
        "source material — never compress a day into a single short sentence such as "
        "'Visit Uplistsikhe and Borjomi.' Instead:\n"
        "  - Every Route Segment (distance/driving time) from the AUTHORITATIVE DAY-WISE "
        "ITINERARY block below must appear as its own activities-array item, e.g. "
        "'Tbilisi → Uplistsikhe: 80 km / 1.5 hrs'.\n"
        "  - Every Description sentence must appear as its own activities-array item — do "
        "not merge multiple sentences into one bullet, and do not drop descriptive details "
        "(names of monuments, historical context, what the traveller does at each stop).\n"
        "  - Any item under 'Optional Activities Mentioned' must appear in the activities "
        "array clearly prefixed as optional, e.g. 'Optional (upon request): Wine tasting at "
        "KTW Winery' — never present it as if it were a guaranteed inclusion, and never "
        "omit it from the itinerary either.\n"
        "  - Preserve the 'Overnight' city exactly as given in the source block.\n\n"


        "## PAX / PERSON / ADULT NORMALIZATION\n"
        "The terms Pax, PAX, Person, Persons, People, Adult, Adults, 'No. of Pax', 'No. of Persons', "
        "'Number of Adults' all refer to the SAME field: number of people travelling.\n"
        "- '2 Pax' = '2 Persons' = '2 Adults' — use whichever the user provided, map to the same count.\n"
        "- If the user provides conflicting values (e.g. '2 Pax' and '3 Adults' in the same message), "
        "ask for clarification before proceeding.\n"
        "- In the JSON 'persons' field, always output as 'N Adults' regardless of which synonym was used.\n\n"

        "## ACCOMMODATION SELECTION — KNOWLEDGE BASE DRIVEN (MANDATORY)\n"
        "This section governs ALL accommodation selection. Follow exactly.\n\n"

        "### A. NEVER HARDCODE\n"
        "Do NOT hardcode countries, cities, hotel names, star categories, or category labels (A/B/C/D).\n"
        "The KB is the ONLY source of truth for hotel names and categories.\n\n"

        "### B. CATEGORY DISCOVERY\n"
        "If the user has NOT specified a hotel category:\n"
        "  1. Search the KB for accommodation data relevant to this destination/country.\n"
        "  2. Extract the available categories (whatever labels the KB uses — '5 Star', '4 Star – C', 'Luxury', etc.).\n"
        "  3. Present those exact categories to the user and ask them to choose.\n"
        "  DO NOT assume or default to any category.\n\n"

        "### C. CATEGORY NORMALIZATION (formatting only)\n"
        "Normalize only formatting differences that refer to the SAME category:\n"
        "  '4 Star – C' = '4 Star C' = '4-Star C' = '4* C' = '4 Star-C'\n"
        "  '5 Star' = '5*' = '5-Star'\n"
        "  Different sub-categories (A, B, C, D) are DISTINCT — never equate them.\n"
        "  DO NOT assume Luxury=5 Star, Premium=4 Star, Standard=3 Star unless KB says so.\n\n"

        "### D. PER-DESTINATION LOOKUP\n"
        "For EVERY destination in the itinerary, independently:\n"
        "  1. Identify the destination name.\n"
        "  2. Find the user's requested accommodation category.\n"
        "  3. Search KB: Country → Destination → Exact Category → Hotel names.\n"
        "  4. The KB may be a table, paragraph, list, heading, or mixed format — understand semantically.\n"
        "  5. Select a hotel ONLY if it satisfies ALL three: correct destination + exact category + present in KB.\n"
        "  6. If multiple hotels exist: pick one suitable option OR list them for user to choose.\n\n"

        "### E. CATEGORY NOT AVAILABLE\n"
        "If the requested category does not exist for a destination in the KB:\n"
        "  Tell the user: '[Category] is not available for [Destination]. Available categories are: [list]. Please select one.'\n"
        "  NEVER silently downgrade, upgrade, or substitute another category.\n\n"

        "### F. HOTEL NAME RULES\n"
        "  - Use EXACT hotel names from the KB. NEVER invent names.\n"
        "  - If KB says 'Hotel A / Hotel B / Similar', pick Hotel A or Hotel B, or list both.\n"
        "  - Preserve the word 'Similar' from KB if no specific hotel was chosen — do not replace with made-up name.\n\n"

        "### G. VALIDATION BEFORE OUTPUT\n"
        "Before inserting any hotel into the itinerary, confirm:\n"
        "  ✓ Hotel belongs to the requested destination\n"
        "  ✓ Hotel belongs to the exact requested category\n"
        "  ✓ Hotel is present in the KB\n"
        "  If any check fails: search KB again. Do not use the hotel.\n\n"

        "### H. BLOCK GENERATION UNTIL CONFIRMED\n"
        "Do NOT generate ---READY--- or ---JSON--- until:\n"
        "  - Hotel category is confirmed by user\n"
        "  - All destinations have a valid hotel from KB\n\n"


        "## RULES\n"
        "- Self-drive inclusion ONLY if itinerary has self-drive\n"
        "- Train tickets ONLY if trains are in this itinerary\n"
        "- Price from KB only; omit field if not available\n"
        "- Do NOT generate ---JSON--- until hotel category is confirmed AND all other selections confirmed\n\n"

        "## INCLUSIONS & EXCLUSIONS SELECTION — PACKAGE-EXACT MATCH (MANDATORY)\n"
        "The '## AUTHORITATIVE INCLUSIONS & EXCLUSIONS' section below contains a SEPARATE "
        "inclusions/exclusions list for EVERY package variant in the knowledge base. Many "
        "packages share the same cities and nearly identical wording (e.g. 'Daily Breakfast', "
        "'Round Airport Transfers', 'Tips for guide and driver' appear in almost every package), "
        "but the NIGHT-BY-NIGHT accommodation breakdown is UNIQUE to each package and MUST match "
        "exactly the itinerary you are building. For example:\n"
        "  - '5 Nights - 4 Nights Tbilisi + 1 Night Gudauri' → inclusions list '3 Nights' "
        "accommodation in Tbilisi', '1 Night accommodation in Gudauri', '1 Night accommodation "
        "in Tbilisi' (i.e. split before/after the Gudauri night)\n"
        "  - '6 Nights - 4 Nights Tbilisi + 2 Nights Batumi' → a COMPLETELY DIFFERENT inclusions "
        "list with Batumi nights instead of Gudauri\n"
        "These must NEVER be mixed, even though most of the other bullet points look identical.\n\n"
        "STEP-BY-STEP:\n"
        "  1. Identify the EXACT package heading that matches the itinerary's total nights and "
        "the exact city + night split (e.g. '5 Nights - 4 Nights Tbilisi + 1 Night Gudauri').\n"
        "  2. Locate the block in '## AUTHORITATIVE INCLUSIONS & EXCLUSIONS' whose 'Package:' "
        "label matches that heading exactly (or as close as possible if the itinerary was "
        "custom-modified from a base package).\n"
        "  3. Copy that block's Inclusions and Exclusions items verbatim into the JSON "
        "'inclusions' and 'exclusions' arrays.\n"
        "  4. Do NOT invent, merge, or borrow any accommodation-night line item from a "
        "different package's block — the night count and city split MUST reconcile with the "
        "itinerary's actual day-by-day overnight cities.\n"
        "  5. If no exact match exists (e.g. a fully custom itinerary), select the CLOSEST "
        "matching package by nights/cities and adapt only the accommodation-night lines to "
        "reflect the actual itinerary, keeping all other generic inclusions/exclusions intact.\n\n"

        "## HOTEL OPTIONS SELECTION — PACKAGE-EXACT MATCH (MANDATORY)\n"
        "The '## AUTHORITATIVE HOTEL OPTIONS' section below contains a SEPARATE hotel-options "
        "table for EVERY package variant in the knowledge base. Match the exact package heading "
        "(same rule as inclusions/exclusions above) before reading hotel names for any city — "
        "then pick the hotel(s) listed for that destination + the user's confirmed star category, "
        "copying the exact hotel names/'or similar' wording verbatim. Do not borrow a hotel name "
        "from a different package's table even if the city name is the same.\n\n"

        "## IMPORTANT NOTES SELECTION — PACKAGE-EXACT MATCH (MANDATORY)\n"
        "The '## AUTHORITATIVE IMPORTANT NOTES' section below contains a SEPARATE Important "
        "Notes block for EVERY package variant in the knowledge base. Match the exact package "
        "heading (same rule as inclusions/exclusions above), then copy that package's notes "
        "verbatim into the JSON 'notes' array — every note in that block must be included; do "
        "not omit, shorten, or merge notes from a different package. The 'notes' array must "
        "never be empty when a matching package's Important Notes block exists in the "
        "knowledge base.\n\n"

        f"## KNOWLEDGE BASE\n{context}\n\n"

        f"## CONVERSATION HISTORY\n{history_text}"
    )

    raw = call_gemini(api_key, model, system, question)

    # Split answer and JSON
    meta = {}
    answer = raw

    # Detect ready separator (---READY--- preferred, ---JSON--- legacy fallback)
    for sep in ("---READY---", "---JSON---"):
        if sep in raw:
            parts = raw.split(sep, 1)
            answer = parts[0].strip()
            json_str = parts[1].strip()
            json_str = re.sub(r'^```[a-z]*\n?', '', json_str).strip()
            json_str = re.sub(r'\n?```$', '', json_str).strip()
            try:
                meta = json.loads(json_str)
            except Exception:
                m = re.search(r'\{.*\}', json_str, re.DOTALL)
                if m:
                    try:
                        meta = json.loads(m.group())
                    except Exception:
                        pass
            break

    if not meta:
        # Fallback: look for a JSON code-block anywhere in the response
        m = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', raw, re.DOTALL)
        if not m:
            m = re.search(r'(\{\s*"destination".*?"notes".*?\})', raw, re.DOTALL)
        if m:
            try:
                meta = json.loads(m.group(1))
                answer = raw[:m.start()].strip()
                if not answer:
                    answer = f"Here is your itinerary for **{meta.get('destination','')}**."
            except Exception:
                pass

    return answer, sources, meta


# ── Paste Content → PDF (analysis, source-restricted, isolated from KB/RAG) ──
# Lets a user paste an itinerary from ANY AI assistant and convert it into a
# KukuTrip PDF WITHOUT touching the Knowledge Base / RAG pipeline at all.

def _extract_json_from_llm_response(raw):
    t = re.sub(r'^```[a-z]*\n?', '', (raw or "").strip())
    t = re.sub(r'\n?```$', '', t).strip()
    try:
        o = json.loads(t)
        if isinstance(o, dict):
            return o
    except Exception:
        pass
    m = re.search(r'\{.*\}', t, re.DOTALL)
    if m:
        try:
            o = json.loads(m.group())
            if isinstance(o, dict):
                return o
        except Exception:
            pass
    return {}

def _validate_pdf_meta(meta):
    if not isinstance(meta, dict) or not meta:
        return False
    days = meta.get("days", [])
    if not isinstance(days, list):
        return False
    for d in days:
        if not isinstance(d, dict):
            return False
        if "activities" in d and d["activities"] is not None and not isinstance(d["activities"], list):
            return False
    for k in ("hotels", "highlights", "inclusions", "exclusions", "notes"):
        if k in meta and meta[k] is not None and not isinstance(meta[k], list):
            return False
    return True

def _normalize_pdf_meta(meta):
    defaults = {"destination": "", "package_name": "", "image_keyword": "", "route": "",
                "dates": "", "nights": "", "persons": "", "transport": "", "days": [],
                "hotels": [], "highlights": [], "inclusions": [], "exclusions": [],
                "notes": [], "amount": ""}
    out = {**defaults, **{k: v for k, v in meta.items() if v is not None}}
    nd = []
    for i, d in enumerate(out.get("days") or [], start=1):
        if not isinstance(d, dict):
            continue
        nd.append({"day": d.get("day", i), "date": d.get("date", "") or "",
                   "title": d.get("title", "") or "",
                   "activities": [a for a in (d.get("activities") or []) if a],
                   "overnight": d.get("overnight", "") or ""})
    out["days"] = nd
    for k in ("hotels", "highlights", "inclusions", "exclusions", "notes"):
        out[k] = [x for x in (out.get(k) or []) if x]
    out["nights"] = _format_nights_days(out.get("nights", ""), out.get("days"))
    if not out.get("package_name"):
        out["package_name"] = out.get("destination") or "Travel Itinerary"
    return out

def analyze_pasted_content(raw_content, api_key, model):
    """Analyze pasted AI-generated content into the existing PDF data model.
    Source-restricted: Knowledge Base/RAG/web search/external knowledge = OFF."""
    system = (
        "You are a Content Analysis & Extraction engine for a travel itinerary PDF "
        "generator. You will be given raw text a user pasted from an AI assistant "
        "(ChatGPT, Claude, Gemini, Copilot, Perplexity, etc). Extract it into a "
        "structured JSON object. This is NOT a generation task — never add travel "
        "knowledge of your own.\n\n"
        "## STRICT SOURCE RESTRICTION\nUse ONLY the pasted content. Knowledge Base=OFF, "
        "RAG=OFF, document retrieval=OFF, web search=OFF, external travel knowledge=OFF. "
        "Do not supplement missing details even if you know them.\n\n"
        "## DO NOT INVENT MISSING VALUES\nIf dates/hotel/price/destination/traveler count "
        "are absent, use empty string \"\" or empty list []. Empty is always better than guessed.\n\n"
        "## CONVERSATIONAL WRAPPER\nExclude AI chat filler ('Absolutely! Here's...', "
        "'Hope you have a great trip!') from extracted fields, but keep factual content.\n\n"
        "## MARKDOWN & TABLES\nUnderstand #, ##, **bold**, lists, and | tables | semantically; "
        "never leave markdown syntax characters in field values.\n\n"
        "## LOSSLESS\nPreserve details that don't fit a specific field inside 'notes' verbatim "
        "rather than dropping them.\n\n"
        "## OUTPUT — JSON SCHEMA (reuse the existing PDF data model exactly)\n"
        "Output ONLY one JSON object, no markdown fences, no explanation:\n"
        "{\n"
        '  "destination": "", "package_name": "", "image_keyword": "", "route": "",\n'
        '  "dates": "", "nights": "", "persons": "", "transport": "",\n'
        '  "days": [{"day":1,"date":"","title":"","activities":[],"overnight":""}],\n'
        '  "hotels": [{"city":"","hotel":"","dates":""}],\n'
        '  "highlights": [], "inclusions": [], "exclusions": [], "notes": [], "amount": ""\n'
        "}\n"
        "'days' MUST be an array (empty if none). Each day's 'activities' MUST be an array "
        "of strings, one item per activity/meal/attraction/transfer — never merge into one "
        "string. Number days sequentially as they appear in the source; never fabricate a day."
    )
    user = f"Analyze and extract the following pasted content:\n\n{raw_content}"
    raw = call_gemini(api_key, model, system, user)
    meta = _extract_json_from_llm_response(raw)
    if not _validate_pdf_meta(meta):
        repair_user = (
            "Your previous output was invalid. Re-extract the SAME content below and output "
            "ONLY one syntactically valid JSON object matching the schema — no fences, no text "
            f"outside the JSON.\n\n{raw_content}"
        )
        raw2 = call_gemini(api_key, model, system, repair_user)
        meta = _extract_json_from_llm_response(raw2)
        if not _validate_pdf_meta(meta):
            raise ValueError("Could not extract structured itinerary data from the pasted content.")
    return _normalize_pdf_meta(meta)

PASTE_MODE_PROMPT_TEMPLATE = """Create a detailed travel itinerary in the following **KukuTrip travel-agent format**.

The output will be copied directly into another application, where the content will be analyzed and converted into a professionally designed PDF.

## IMPORTANT OUTPUT RULES

* Follow the structure below closely.
* Do NOT explain the format.
* Do NOT add conversational introductions or conclusions.
* Do NOT use JSON.
* Do NOT use XML.
* Do NOT provide code.
* Do NOT include comments about how the itinerary was created.
* Use clean headings, numbered days, dates, and bullet points.
* Keep the output suitable for direct copy/paste into a travel itinerary PDF generator.
* Use factual information provided in the request.
* If a detail is not provided, do not invent it.
* Clearly mark estimates or assumptions if they are explicitly requested.
* Keep the itinerary professional and travel-agent friendly.

---

# REQUIRED FORMAT

## Trip Header

Start with:

**[START DATE] – [END DATE] [DURATION] [TRAVELER COUNT] [TRANSPORT TYPE]**

Example:

05 Oct – 14 Oct 2026
9 Nights / 10 Days
2 Adults · 0 Children · 1 Room
Private Transfers + High-Speed Trains

Then:

**Day-wise Itinerary**

**[TRIP TITLE IN UPPERCASE]**

[City 1] · [City 2] · [City 3]

[START DATE] – [END DATE]

---

# DAY-WISE FORMAT

For every day use this structure:

### [DAY NUMBER]

## [DAY TITLE]

[DATE]

• [Activity / detail]
• [Activity / detail]
• [Activity / detail]
• [Activity / detail]
• [Activity / detail]
• [Activity / detail]
• Overnight in [City]

Each day should contain meaningful details rather than a short one-line summary.

Include, where applicable:

* Breakfast
* Departure time
* Check-out
* Check-in
* Driving distance
* Approximate driving time
* Train journey
* Airport transfer
* Sightseeing
* Attractions
* Activities
* Meals
* Free time
* Shopping
* Hotel
* Overnight location

---

# EXAMPLE DAY

### 1

## Arrival in Rome

05 OCT

• Arrival at Rome Fiumicino Airport.
• Meet the driver at the arrivals area.
• Transfer to the hotel for check-in.
• Standard check-in time – 15:00.
• Free time to freshen up and relax.
• Evening visit to Piazza Navona.
• Visit the Trevi Fountain.
• Enjoy dinner at a local Italian restaurant.
• Overnight in Rome.

---

# HOTELS SECTION

After the day-wise itinerary, add:

## Hotels Confirmed

Use a table:

| City / Nights       | Hotel             | Dates   | Meal Plan       |
| -------------------- | ------------------ | ------- | ---------------- |
| Rome (3 Nights)     | [Hotel] / Similar | [Dates] | Bed & Breakfast |
| Florence (2 Nights) | [Hotel] / Similar | [Dates] | Bed & Breakfast |

If hotels are not provided, use:

**To be confirmed**

Do not invent hotel names.

---

# TOTAL PACKAGE PRICE

Add:

## TOTAL PACKAGE PRICE

[Adults] · [Children] · [Rooms] · [Transport]

[Currency + Amount]

If the user has not provided a price, write:

**Price to be confirmed**

Do not invent a package price.

---

# PACKAGE HIGHLIGHTS

Add:

## Package Highlights

Provide 4–8 concise highlights based ONLY on the itinerary.

Example:

• Colosseum & Ancient Rome
• Vatican Museums & Sistine Chapel
• Florence Renaissance Experience
• Venice Gondola Ride
• Lake Como Excursion

Do not introduce activities that are not included in the itinerary.

---

# INCLUSIONS

Add:

## ✓ Inclusions

Include only services actually mentioned or explicitly included in the request.

Examples:

• Accommodation as specified
• Daily breakfast
• Airport transfers
• Private transportation
• High-speed train tickets
• English-speaking guide
• Sightseeing excursions
• Entrance tickets specifically mentioned
• Local taxes and charges

Do not automatically assume something is included.

---

# EXCLUSIONS

Add:

## ✕ Exclusions

Include exclusions provided by the user.

If the user has not provided explicit exclusions, use only reasonable generic categories when requested, such as:

• International flights
• Travel insurance
• Personal expenses
• Tips
• Meals not mentioned in the itinerary
• Services not mentioned in the program

Do not state that a service is excluded if the user explicitly said it is included.

---

# IMPORTANT NOTES

Add:

## ⚑ Important Notes

Include practical information relevant to the itinerary, but do NOT introduce unsupported factual claims.

Examples:

• Hotel check-in and check-out times are subject to hotel policy.
• Comfortable walking shoes are recommended.
• Attraction timings are subject to availability.
• Train timings are subject to availability.
• Any additional sightseeing not mentioned in the itinerary may incur additional costs.
• Any increase in applicable taxes or transportation costs may be recovered from the guest.

Only include commercial/payment conditions if they were provided by the user or explicitly requested.

---

# CONTENT STYLE

The final output should look like a professional travel company's itinerary document.

Use:

* Clear section hierarchy
* Short paragraphs
* Bullet points beginning with `•`
* Numbered days
* Dates in `DD MMM`
* City names consistently
* Professional travel terminology
* Concise but sufficiently detailed descriptions

Avoid:

* Long conversational paragraphs
* Marketing-heavy language
* Excessive emojis
* "Here is your itinerary..."
* "I hope you enjoy..."
* AI commentary
* Markdown code blocks
* JSON
* Explanations outside the itinerary

---

# DATA INTEGRITY

The itinerary will later be processed automatically by a PDF generator.

Therefore:

1. Never invent hotel names.
2. Never invent prices.
3. Never invent flight details.
4. Never invent travel dates.
5. Never invent traveler counts.
6. Never invent included services.
7. Never invent attraction tickets.
8. Never add activities that were not requested.
9. If information is unavailable, write **To be confirmed** where appropriate.
10. Preserve all information supplied by the user.

---

# INPUT

Create the itinerary using the following trip information:

**Destination:** [DESTINATION]

**Travel Dates:** [START DATE] – [END DATE]

**Travelers:** [ADULTS] Adults · [CHILDREN] Children · [ROOMS] Rooms

**Transportation:** [TRANSPORTATION]

**Cities / Route:** [ROUTE]

**Hotels:** [HOTELS, IF AVAILABLE]

**Budget / Package Price:** [PRICE, IF AVAILABLE]

**Special Requirements:** [SPECIAL REQUIREMENTS]

**Places / Activities Requested:**
[LIST OF REQUESTED PLACES AND ACTIVITIES]

**Inclusions:**
[INCLUSIONS]

**Exclusions:**
[EXCLUSIONS]

**Important Notes:**
[IMPORTANT NOTES]

Now generate the complete itinerary using the KukuTrip format above.
"""


def render_paste_content_mode(api_key, model, selected_theme_name):
    """Isolated UI for Paste Content → PDF. Never touches Knowledge Base/RAG."""
    st.title("✈ Travel Itinerary Agent")
    st.caption(
        "Paste content directly from ChatGPT, Claude, Gemini, or any other AI "
        "assistant. The content will be analyzed and converted into the KukuTrip "
        "PDF format automatically."
    )
    with st.expander("📋 Copy prompt template (use this in ChatGPT/Claude/Gemini first)"):
        st.caption(
            "Copy this prompt into your AI assistant of choice, fill in the trip "
            "details, and paste the AI's response into the box below to generate "
            "a KukuTrip-formatted PDF."
        )
        st.code(PASTE_MODE_PROMPT_TEMPLATE, language=None)

    raw_content = st.text_area("Paste your content", height=350, key="paste_raw_content",
                                placeholder="Paste your AI-generated itinerary here…")

    price_basis_choice = st.radio(
        "Price Basis",
        ["Auto-detect from content", "Per Person", "Total Package"],
        key="paste_price_basis_choice",
        horizontal=True,
        help="The price amount is taken exactly as written in the pasted content "
             "(no calculation is performed). This only controls whether it is "
             "labeled 'Price Per Person' or 'Total Package Cost' on the PDF.",
    )
    st.caption(f"🎨 Using PDF Theme: **{selected_theme_name}** (change in sidebar)")

    if st.button("📄 Generate PDF", type="primary", use_container_width=True, key="paste_generate_pdf_btn"):
        if not raw_content or not raw_content.strip():
            st.error("Please paste some content before generating the PDF.")
            return
        with st.spinner("Analyzing pasted content…"):
            try:
                structured = analyze_pasted_content(raw_content, api_key, model)
            except Exception as e:
                st.error(f"Could not analyze the pasted content: {e}")
                return

        # ── Price: use the extracted amount exactly as provided — no
        # calculation/conversion/markup is applied. The basis choice only
        # selects the display label (Per Person vs Total Package).
        raw_amount = (structured.get("amount") or "").strip()
        if raw_amount:
            if price_basis_choice == "Per Person":
                is_pp = True
            elif price_basis_choice == "Total Package":
                is_pp = False
            else:
                is_pp = _is_per_person_amount(raw_amount)
            structured["_final_price_label"] = "PRICE PER PERSON" if is_pp else "TOTAL PACKAGE COST"
            structured["_final_price_value"] = raw_amount

        with st.spinner("Generating PDF…"):
            hdr_img = None
            try:
                img_kw = structured.get("image_keyword") or structured.get("destination") or ""
                if img_kw:
                    hdr_img = fetch_destination_image(img_kw)
            except Exception:
                hdr_img = None
            try:
                pdf_bytes = generate_pdf_themed(
                    structured.get("package_name") or "Itinerary",
                    structured, selected_theme_name, header_img_path=hdr_img,
                )
            except Exception as e:
                st.error(f"PDF generation failed: {e}")
                return
        st.session_state["paste_pdf_bytes"] = pdf_bytes
        st.session_state["paste_pdf_name"] = f"itinerary_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
        st.session_state["paste_pdf_structured"] = structured

    if st.session_state.get("paste_pdf_bytes"):
        st.success("✅ PDF generated successfully!")
        st.download_button(
            "⬇️ Download Itinerary PDF", data=st.session_state["paste_pdf_bytes"],
            file_name=st.session_state.get("paste_pdf_name", "itinerary.pdf"),
            mime="application/pdf", key="paste_download_pdf",
        )
        with st.expander("🔍 Extracted structured data (for review)"):
            st.json(st.session_state.get("paste_pdf_structured", {}))


# ── Sidebar ───────────────────────────────────────────────────────────────────

with st.sidebar:
    st.title("✈ Travel Agent")
    st.caption("Powered by Gemini + ChromaDB")

    _default_key = (
        st.secrets.get("GOOGLE_API_KEY", "")
        if hasattr(st, "secrets") else ""
    ) or os.environ.get("GOOGLE_API_KEY", "")
    api_key = st.text_input(
        "Google Gemini API Key", type="password",
        value=_default_key,
        help="Free key from https://aistudio.google.com — or set GOOGLE_API_KEY in Streamlit secrets.",
    )
    model = st.selectbox(
        "Model",
        ["gemini-3.1-flash-lite", "gemini-flash-latest", "gemini-3.6-flash"],
        help="gemini-3.1-flash-lite is free.",
    )

    st.divider()
    st.subheader("🎨 PDF Settings")
    pdf_theme_name = st.selectbox(
        "PDF Theme",
        list(PDF_THEMES.keys()),
        index=list(PDF_THEMES.keys()).index(DEFAULT_PDF_THEME),
        help="'Maroon Red' is the classic KukuTrip design. 'Modern Editorial' "
             "uses a clean document-style layout with large day numbers. "
             "All other themes reuse the classic layout with a different "
             "color palette.",
    )
    _theme_preview = PDF_THEMES[pdf_theme_name]
    st.markdown(
        f"<div style='display:flex;gap:6px;align-items:center;margin-top:-6px;'>"
        f"<span style='width:14px;height:14px;border-radius:50%;background:{_theme_preview['primary']};display:inline-block;border:1px solid #ccc;'></span>"
        f"<span style='width:14px;height:14px;border-radius:50%;background:{_theme_preview['accent']};display:inline-block;border:1px solid #ccc;'></span>"
        f"<span style='width:14px;height:14px;border-radius:50%;background:{_theme_preview['bg_light']};display:inline-block;border:1px solid #ccc;'></span>"
        f"<span style='font-size:11px;color:#777;'>Primary · Accent · Background</span></div>",
        unsafe_allow_html=True,
    )
    st.session_state["pdf_theme_name"] = pdf_theme_name

    st.divider()
    st.subheader("⚙️ PDF Generation Mode")
    generation_mode = st.radio(
        "Mode",
        ["Knowledge Base", "Paste Content → PDF"],
        key="generation_mode",
        help="'Knowledge Base' uses your uploaded PDF/Word documents (RAG) or "
             "the KukuTrip Master Travel Plans Excel database, per the "
             "'Knowledge Source' selector below. 'Paste Content → PDF' "
             "converts content pasted from ChatGPT, Claude, Gemini, or any "
             "AI assistant directly into a PDF — no knowledge source is "
             "used in this mode.",
    )
    # Only relevant in Knowledge Base mode — declared here so `existing`
    # always exists (as an empty default) regardless of which mode is active.
    existing = []

    if generation_mode == "Knowledge Base":
        st.divider()
        st.subheader("🗂️ Knowledge Source")
        knowledge_source = st.selectbox(
            "Knowledge Source",
            [KB_SOURCE_EXISTING, KB_SOURCE_MASTER_PLANS],
            key="knowledge_source",
            help=f"'{KB_SOURCE_EXISTING}' uses your uploaded PDF/Word documents "
                 f"(unchanged RAG pipeline). '{KB_SOURCE_MASTER_PLANS}' uses the "
                 f"pre-built KukuTrip package database "
                 f"({MASTER_PLANS_XLSX.name}) instead — these two sources are "
                 "completely independent; switching between them does not "
                 "affect the other's data or indexing.",
        )
        if knowledge_source == KB_SOURCE_MASTER_PLANS:
            if MASTER_PLANS_XLSX.exists():
                st.caption(f"✅ Found `{MASTER_PLANS_XLSX.name}` under docs/")
            else:
                st.warning(
                    f"⚠️ `{MASTER_PLANS_XLSX.name}` not found under `docs/`. "
                    "Please add the file to use this source."
                )
    else:
        knowledge_source = KB_SOURCE_EXISTING

    if generation_mode == "Knowledge Base" and knowledge_source == KB_SOURCE_EXISTING:
        st.divider()
        st.subheader("📂 Travel Documents")

        uploaded = st.file_uploader("Upload PDF or Word files", type=["pdf", "docx"], accept_multiple_files=True)
        if uploaded:
            DOCS_DIR.mkdir(exist_ok=True)
            for f in uploaded:
                (DOCS_DIR / f.name).write_bytes(f.read())
            st.success(f"Saved {len(uploaded)} file(s)")
            if CHROMA_DIR.exists():
                shutil.rmtree(CHROMA_DIR)
            st.cache_resource.clear()
            st.session_state.pop("messages", None)

        existing = (list(DOCS_DIR.glob("*.pdf")) + list(DOCS_DIR.glob("*.docx"))) if DOCS_DIR.exists() else []
        if existing:
            st.divider()
            st.subheader("📄 Indexed Documents")
            for f in sorted(existing):
                st.write(f"• {f.name}")

        st.divider()
        if st.button("🔄 Re-index Documents", use_container_width=True):
            if CHROMA_DIR.exists():
                shutil.rmtree(CHROMA_DIR)
            st.cache_resource.clear()
            st.session_state.pop("messages", None)
            st.rerun()

        # History
        st.divider()
        st.subheader("🕘 Previous Chats")
        st.caption(f"Saved for {RETENTION} days")
        col_new, col_clear = st.columns(2)
        with col_new:
            if st.button("➕ New Chat", use_container_width=True):
                st.session_state.session_id = new_sid()
                st.session_state.messages = []
                st.rerun()
        with col_clear:
            if st.button("🗑️ Clear All", use_container_width=True):
                if HISTORY_DIR.exists():
                    shutil.rmtree(HISTORY_DIR)
                HISTORY_DIR.mkdir(exist_ok=True)
                st.session_state.session_id = new_sid()
                st.session_state.messages = []
                st.rerun()

        for sf in list_sessions()[:15]:
            label = sf.stem.replace("_", " ", 1)
            if st.button(f"📂 {label}", key=str(sf), use_container_width=True):
                st.session_state.messages = load_session(sf)
                st.session_state.session_id = sf.stem
                st.rerun()


# ── Main ──────────────────────────────────────────────────────────────────────

if not api_key:
    st.title("✈ Travel Itinerary Agent")
    st.warning("Enter your **Google Gemini API key** in the sidebar. Free at https://aistudio.google.com")
    st.stop()

# ── Isolated code path: Paste Content → PDF ──────────────────────────────────
# This branch NEVER touches the Knowledge Base / RAG pipeline below it.
if st.session_state.get("generation_mode") == "Paste Content → PDF":
    render_paste_content_mode(
        api_key, model,
        st.session_state.get("pdf_theme_name", DEFAULT_PDF_THEME),
    )
    st.stop()

# ── Isolated code path: KukuTrip Master Travel Plans ─────────────────────────
# Completely separate from the existing ChromaDB/docx knowledge base below —
# never calls build_vectorstore(), never touches CHROMA_DIR, never reads
# DOCS_DIR's pdf/docx files. Reuses the same pending_meta/pdf_ready/
# "Itinerary Ready Card" pricing+PDF-generation flow at the bottom of this
# file, so the existing PDF generator is reused unchanged for both sources.
if st.session_state.get("knowledge_source") == KB_SOURCE_MASTER_PLANS:
    st.title("✈ Travel Itinerary Agent")
    st.caption(
        "Knowledge Source: **KukuTrip Master Travel Plans** — ask with your "
        "destination, dates & travelers to get a formatted itinerary + PDF."
    )

    _mp_sheets = get_master_plans_workbook()
    if _mp_sheets is None:
        st.error(
            f"⚠️ Could not load `{MASTER_PLANS_XLSX.name}` from `docs/`. "
            "Please make sure the file exists and is a valid Excel workbook."
        )
        st.stop()
    st.success(
        f"✅ Loaded **{len(_mp_sheets.get('Packages', []))}** packages from "
        f"KukuTrip Master Travel Plans. Ready!"
    )

    if "mp_messages" not in st.session_state:
        st.session_state.mp_messages = [
            {"role": "assistant", "content":
             "Hi! I'm your KukuTrip travel consultant. Tell me your destination, "
             "travel dates, and number of travelers — I'll find the best matching "
             "package from our master database and build your itinerary! 🗺️"}
        ]

    _mp_latest_pdf_idx = -1
    for _i, _m in enumerate(st.session_state.mp_messages):
        if _m.get("pdf"):
            _mp_latest_pdf_idx = _i

    for idx, msg in enumerate(st.session_state.mp_messages):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg.get("pdf") and idx == _mp_latest_pdf_idx:
                st.download_button(
                    "⬇️ Download Itinerary PDF",
                    data=bytes.fromhex(msg["pdf"]),
                    file_name=msg.get("pdf_name", "itinerary.pdf"),
                    mime="application/pdf",
                    key=f"mp_dl_{msg.get('_id', id(msg))}",
                )

    if mp_prompt := st.chat_input(
        "e.g. I want to visit Georgia for 6 nights with 2 adults",
        key="mp_chat_input",
    ):
        st.session_state.mp_messages.append({"role": "user", "content": mp_prompt})
        with st.chat_message("user"):
            st.markdown(mp_prompt)

        with st.chat_message("assistant"):
            with st.spinner("Searching KukuTrip Master Travel Plans…"):
                try:
                    hints = extract_master_plans_hints(mp_prompt)
                    dest_hint = guess_master_plans_destination(_mp_sheets, mp_prompt)
                    answer, sources, meta = get_master_plans_answer(
                        mp_prompt, st.session_state.mp_messages[:-1], api_key, model,
                        destination_hint=dest_hint,
                        nights_hint=hints["nights"], days_hint=hints["days"],
                    )
                    # NOTE: matched Package IDs are intentionally NOT appended to
                    # the chat answer — showing that block on every reply pushed
                    # the "Generate PDF" button out of view / below the fold.
                except Exception as e:
                    answer = f"❌ Error: {e}"
                    meta = {}

            if meta and not answer.startswith("❌"):
                st.session_state.pending_meta = meta
                st.session_state.pdf_ready = True
                for _k in ("price_conv_rate", "price_display_mode", "price_markup_raw",
                           "price_workflow_done", "final_price_label", "final_price_value"):
                    st.session_state.pop(_k, None)
            elif not answer.startswith("❌"):
                st.session_state.pdf_ready = False

            msg_entry = {
                "role": "assistant",
                "content": answer,
                "_id": len(st.session_state.mp_messages),
            }
            if meta:
                msg_entry["meta"] = meta
            st.session_state.mp_messages.append(msg_entry)
            st.rerun()

# ── Knowledge Base mode (existing RAG flow, unchanged) ───────────────────────
else:
    st.title("✈ Travel Itinerary Agent")
    st.caption("Ask with dates & travelers — get a formatted itinerary + PDF instantly.")

    if not existing:
        st.info("Upload travel documents (PDF/DOCX) in the sidebar to get started.")
        st.stop()

    with st.spinner("Loading document index…"):
        try:
            col = build_vectorstore()
        except Exception as e:
            st.error(f"Failed to build index: {e}")
            st.stop()

    if col is None:
        st.info("No documents indexed yet — upload files in the sidebar.")
        st.stop()

    # col.count() can fail if persisted DB schema is stale — auto-rebuild if so
    try:
        chunk_count = col.count()
    except Exception:
        # Stale/incompatible DB — wipe and rebuild
        if CHROMA_DIR.exists():
            shutil.rmtree(CHROMA_DIR)
        st.cache_resource.clear()
        st.rerun()

    st.success(f"✅ Indexed **{chunk_count}** chunks from {len(existing)} file(s). Ready!")

    if "messages" not in st.session_state:
        st.session_state.messages = [
            {"role": "assistant", "content":
             "Hi! Tell me your travel dates, number of travelers, and destination — "
             "I'll create a personalised itinerary using your documents and generate a PDF! 🗺️"}
        ]

    # Find the index of the most recent message with a PDF
    _latest_pdf_idx = -1
    for _i, _m in enumerate(st.session_state.messages):
        if _m.get("pdf"):
            _latest_pdf_idx = _i

    for idx, msg in enumerate(st.session_state.messages):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            # Only show download button on the most recent PDF message
            if msg.get("pdf") and idx == _latest_pdf_idx:
                st.download_button(
                    "⬇️ Download Itinerary PDF",
                    data=bytes.fromhex(msg["pdf"]),
                    file_name=msg.get("pdf_name", "itinerary.pdf"),
                    mime="application/pdf",
                    key=f"dl_{msg.get('_id', id(msg))}",
                )

    if prompt := st.chat_input("e.g. Plan a 5-day Georgia trip for 2 adults + 1 kid, 20-25 Aug 2026"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Gemini is crafting your itinerary…"):
                try:
                    answer, sources, meta = get_answer(
                        col, prompt, st.session_state.messages[:-1], api_key, model
                    )
                    if sources:
                        answer += f"\n\n---\n*📄 Sources: {', '.join(sorted(sources))}*"
                except Exception as e:
                    answer = f"❌ Error: {e}"
                    meta = {}

            # When meta present, store as pending — do NOT auto-generate PDF
            if meta and not answer.startswith("❌"):
                st.session_state.pending_meta = meta
                st.session_state.pdf_ready = True
                # Reset pricing workflow state for the new itinerary
                for _k in ("price_conv_rate", "price_display_mode", "price_markup_raw",
                           "price_workflow_done", "final_price_label", "final_price_value"):
                    st.session_state.pop(_k, None)
            elif not answer.startswith("❌"):
                # A clarification/question reply — itinerary not yet ready
                st.session_state.pdf_ready = False

            msg_entry = {
                "role": "assistant",
                "content": answer,
                "_id": len(st.session_state.messages),
            }
            if meta:
                msg_entry["meta"] = meta

            st.session_state.messages.append(msg_entry)
            save_session(st.session_state.session_id, st.session_state.messages)
            st.rerun()


# ── Itinerary Ready Card + Pricing Workflow + Generate PDF Button ─────────────
if st.session_state.get("pdf_ready") and st.session_state.get("pending_meta"):
    meta_stored = st.session_state.pending_meta
    pkg_label = meta_stored.get("package_name") or meta_stored.get("destination") or "Itinerary"

    raw_amount = meta_stored.get("amount", "")
    base_value = _extract_amount_value(raw_amount)
    is_usd = _is_usd_amount(raw_amount)
    is_per_person = _is_per_person_amount(raw_amount) if raw_amount else None
    persons_count = _extract_persons_count(meta_stored.get("persons", ""))


    with st.container(border=True):
        st.success(f"✅ **Itinerary Ready** — {pkg_label}")
        st.caption("Complete the pricing steps below, then generate your PDF.")

        # Show persistent image fetch debug info (survives rerun)
        if st.session_state.get("_img_dbg_msgs"):
            for _m in st.session_state["_img_dbg_msgs"]:
                st.caption(f"🖼️ {_m}")

        # Show the original base price (as found in the KB) for the agent's
        # reference only — this is NEVER shown in the customer-facing PDF.
        if raw_amount:
            st.caption(f"💰 Base price from knowledge base (internal use only): **{raw_amount}**")

        if base_value is None:
            st.warning(
                "⚠️ No base package price was found in the knowledge base for this itinerary. "
                "Please provide a base price manually to continue."
            )

        # ── Manual base price override — always available, with INR/USD toggle ──
        # Lets the agent overwrite the knowledge-base price (or supply one when
        # none was found) in either currency. If left blank, the default base
        # price picked from the knowledge base (shown above) is used as-is.
        # When INR is chosen (default or override), the USD→INR conversion
        # step is skipped entirely; when USD is chosen, conversion is required.
        with st.expander(
            "✏️ Override base price (optional)" if base_value is not None else "✏️ Enter base price",
            expanded=(base_value is None),
        ):
            if base_value is not None:
                st.caption(
                    "Leave the field below blank to keep using the default base price "
                    f"from the knowledge base ({raw_amount})."
                )
            override_currency = st.radio(
                "Currency of the price you're entering",
                ["USD", "INR"],
                index=0 if is_usd else 1,   # default to the currency already detected from the KB
                key="price_override_currency_input",
                horizontal=True,
            )
            manual_amount = st.text_input(
                f"Base package price in {override_currency} (numbers only, no currency symbol) — "
                "leave blank to use the default",
                key="price_manual_base",
            )
            manual_val = _extract_amount_value(manual_amount)
            if manual_val:
                # User explicitly overrode the price — use their value + currency.
                base_value = manual_val
                is_usd = (override_currency == "USD")
            # else: no override entered — base_value/is_usd remain the KB defaults



        # ── Confirm price basis (per-person vs total package) — always shown,
        # defaulted from detection but user-confirmable to avoid ambiguity. ──
        if base_value is not None:
            basis_default_idx = 0 if is_per_person else 1
            basis_choice = st.radio(
                "Confirm: is the base price above per person or total package price?",
                ["Per Person", "Total Package"],
                index=basis_default_idx,
                key="price_basis_confirm_input",
                horizontal=True,
            )
            is_per_person = (basis_choice == "Per Person")

        conv_rate = None
        if base_value is not None and is_usd:
            st.markdown("**Step 1 — USD → INR Conversion Rate**")
            conv_rate_raw = st.text_input(
                "What is the USD → INR conversion rate to use?",
                key="price_conv_rate_input",
                placeholder="e.g. 84.5",
            )
            conv_rate = _extract_amount_value(conv_rate_raw)
            if conv_rate_raw and not conv_rate:
                st.error("Please enter a valid numeric conversion rate.")

        # Base price in INR (after conversion if needed) — remains on the
        # SAME basis (per-person or total) as the original amount; this basis
        # carries through agent markup and markup, and is only resolved to
        # the user's chosen display format in the final step.
        base_inr = None
        if base_value is not None:
            if is_usd:
                if conv_rate:
                    base_inr = base_value * conv_rate
            else:
                base_inr = base_value


        agent_markup_raw = None
        if base_inr is not None:
            st.markdown("**Step 2 — Agent Markup (internal, percentage only)**")
            agent_markup_raw = st.text_input(
                "What is your agent markup percentage, applied on the converted total? "
                "(internal use only — never shown to the customer or in the PDF)",
                key="price_agent_markup_input",
                placeholder="e.g. 10%",
            )

        # ── Compute base (after conversion) + agent markup (internal only) ──
        base_after_agent_markup = None
        agent_markup_val = None
        if base_inr is not None and agent_markup_raw:
            m = re.search(r'[\d.]+', agent_markup_raw)
            agent_pct = float(m.group()) if m else 0.0
            agent_markup_val = base_inr * (agent_pct / 100.0)
            base_after_agent_markup = base_inr + agent_markup_val
            st.caption(
                f"🧮 Internal calculation (not shown to customer): "
                f"Base ₹{_format_inr(base_inr)} + Agent Markup {agent_pct}% "
                f"(₹{_format_inr(agent_markup_val)}) = ₹{_format_inr(base_after_agent_markup)}"
            )

        markup_raw = None
        if base_after_agent_markup is not None:
            st.markdown("**Step 3 — Markup**")
            markup_raw = st.text_input(
                "What markup would you like to add to the package price? "
                "(percentage e.g. 15% or fixed amount e.g. 20000)",
                key="price_markup_input",
                placeholder="e.g. 15% or 20000",
            )

        # ── Total after customer markup (internal, before display split) ────
        total_final = None
        if base_after_agent_markup is not None and markup_raw:
            markup_val = _parse_markup(markup_raw, base_after_agent_markup)
            total_final = base_after_agent_markup + markup_val

        display_mode = None
        if total_final is not None:
            st.markdown("**Step 4 — Price Display in PDF**")
            display_mode = st.radio(
                "How would you like the price to be shown in the PDF?",
                ["Price Per Person", "Total Package Price"],
                key="price_display_mode_input",
                horizontal=True,
            )

        # ── Compute final customer-facing price ─────────────────────────────
        # `total_final` is on the SAME basis as the original base amount
        # (per-person or total package), tracked via `is_per_person`.
        # Convert to whichever basis the user wants displayed, without
        # ever double-dividing/multiplying by persons_count.
        final_label = None
        final_value_str = None
        ready_to_generate = False

        if total_final is not None and display_mode:
            if is_per_person:
                per_person_amt = total_final
                total_amt = total_final * max(1, persons_count)
            else:
                total_amt = total_final
                per_person_amt = total_final / max(1, persons_count)

            if display_mode == "Price Per Person":
                final_label = "Price Per Person"
                final_value_str = f"₹{_format_inr(per_person_amt)} / person"
            else:
                final_label = "Total Package Price"
                final_value_str = f"₹{_format_inr(total_amt)}"
            ready_to_generate = True


        if ready_to_generate:
            st.info(f"**{final_label}:** {final_value_str}")

        gen_disabled = not ready_to_generate
        if gen_disabled:
            st.caption(
                "⏳ Complete all pricing steps above (conversion rate if applicable, "
                "price display option, and markup) before generating the PDF."
            )

        if st.button("📄 Generate PDF", type="primary", use_container_width=True, disabled=gen_disabled):
            with st.spinner("Generating PDF — fetching destination image…"):
                try:
                    # Inject only the final computed price into meta — never the
                    # base/USD/markup/conversion details, per pricing policy.
                    meta_for_pdf = dict(meta_stored)
                    meta_for_pdf["_final_price_label"] = final_label
                    meta_for_pdf["_final_price_value"] = final_value_str

                    img_kw  = meta_stored.get("image_keyword") or meta_stored.get("destination") or "travel landscape"
                    _img_dbg: list = []
                    hdr_img = fetch_destination_image(img_kw, _dbg=_img_dbg)
                    # If specific keyword fails, try just the destination country
                    if hdr_img is None:
                        fallback_kw = meta_stored.get("destination", "travel landscape")
                        if fallback_kw.lower() != img_kw.lower():
                            hdr_img = fetch_destination_image(fallback_kw, _dbg=_img_dbg)
                    st.session_state["_img_dbg_msgs"] = _img_dbg  # persist across rerun
                    selected_theme_name = st.session_state.get("pdf_theme_name", DEFAULT_PDF_THEME)
                    pdf_b = generate_pdf_themed(
                        pkg_label, meta_for_pdf, selected_theme_name,
                        header_img_path=hdr_img,
                    )
                    pdf_n = f"itinerary_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"

                    # Save PDF to template dir
                    TEMPLATE_DIR.mkdir(exist_ok=True)
                    (TEMPLATE_DIR / pdf_n).write_bytes(pdf_b)

                    # Attach the generated PDF to the last assistant message in
                    # whichever chat history list is currently active — the
                    # existing Knowledge Base flow uses `messages`, while the
                    # isolated KukuTrip Master Travel Plans flow uses its own
                    # `mp_messages` list. Attaching to the wrong (empty/unused)
                    # list means no rendered chat message ever gets the "pdf"
                    # key set, so the Download button never appears even
                    # though the PDF was generated successfully.
                    _active_msgs_key = (
                        "mp_messages"
                        if st.session_state.get("knowledge_source") == KB_SOURCE_MASTER_PLANS
                        else "messages"
                    )
                    _active_msgs = st.session_state.get(_active_msgs_key, [])
                    for i in range(len(_active_msgs) - 1, -1, -1):
                        if _active_msgs[i].get("role") == "assistant":
                            _active_msgs[i]["pdf"]      = pdf_b.hex()
                            _active_msgs[i]["pdf_name"] = pdf_n
                            break
                    # Append image debug info to last assistant message for visibility
                    if _img_dbg:
                        img_note = "\n\n*🖼️ Image: " + " | ".join(_img_dbg) + "*"
                        for i in range(len(_active_msgs) - 1, -1, -1):
                            if _active_msgs[i].get("role") == "assistant":
                                _active_msgs[i]["content"] += img_note
                                break
                    st.session_state[_active_msgs_key] = _active_msgs
                    # Session history (save_session) only applies to the
                    # Knowledge Base chat — the Master Travel Plans chat isn't
                    # persisted to disk history, matching its existing behavior.
                    if _active_msgs_key == "messages":
                        save_session(st.session_state.session_id, st.session_state.messages)
                    st.session_state.pdf_ready = False
                    st.session_state.last_pdf       = pdf_b
                    st.session_state.last_pdf_name  = pdf_n
                    st.rerun()
                except Exception as e:
                    st.error(f"PDF generation failed: {e}")

